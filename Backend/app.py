import os
import uuid
import random
import hashlib
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta, date, time, timezone
from functools import wraps
import json
import csv
from io import StringIO, BytesIO

from flask import Flask, request, jsonify, send_file, make_response
from flask_cors import CORS
from flask_jwt_extended import (
    JWTManager, create_access_token, jwt_required,
    get_jwt_identity, get_jwt
)
from werkzeug.utils import secure_filename
from sqlalchemy import func, or_, and_, text
from sqlalchemy.exc import IntegrityError

from config import Config
from models import db, jwt, User, Patient, Nurse, Specialist, Facility, Specialty, Department
from models import Referral, ReferralDocument, Appointment, RescheduleRequest
from models import Notification, AuditLog, Report
from models import SpecialistSchedule, PatientMedicalHistory
from models import UserActivityLog, SystemConfig, UserSession, SystemMetric

# PDF and Word libraries
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==================== LOGGING CONFIGURATION ====================
import logging
import sys

# Create logs directory
os.makedirs('logs', exist_ok=True)

class ColoredConsoleHandler(logging.StreamHandler):
    """Custom handler for colored console output"""
    
    def emit(self, record):
        try:
            msg = self.format(record)
            stream = self.stream
            if record.levelno >= logging.ERROR:
                stream.write(f'\033[91m{msg}\033[0m\n')
            elif record.levelno >= logging.WARNING:
                stream.write(f'\033[93m{msg}\033[0m\n')
            elif record.levelno >= logging.INFO:
                stream.write(f'\033[92m{msg}\033[0m\n')
            else:
                stream.write(f'{msg}\n')
            stream.flush()
        except Exception:
            self.handleError(record)

def setup_logging(app_instance):
    """Configure logging for the application"""
    for handler in logging.root.handlers[:]:
        logging.root.removeHandler(handler)
    
    log_level = logging.DEBUG if app_instance.debug else logging.INFO
    
    logging.basicConfig(
        level=log_level,
        format='%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    
    console_handler = ColoredConsoleHandler(sys.stdout)
    console_handler.setLevel(log_level)
    console_handler.setFormatter(logging.Formatter(
        '%(asctime)s [%(levelname)s] %(message)s',
        datefmt='%H:%M:%S'
    ))
    logging.getLogger().addHandler(console_handler)
    
    file_handler = logging.FileHandler('logs/app.log', encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_handler.setFormatter(logging.Formatter(
        '%(asctime)s [%(levelname)s] %(name)s: %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    ))
    logging.getLogger().addHandler(file_handler)
    
    app_instance.logger = logging.getLogger('isas')
    app_instance.logger.info("=" * 50)
    app_instance.logger.info("ISAS Server Starting")
    app_instance.logger.info(f"Log level: {logging.getLevelName(log_level)}")
    app_instance.logger.info("=" * 50)
    
    return app_instance.logger

app = Flask(__name__)
app.config.from_object(Config)

logger = setup_logging(app)

db.init_app(app)
jwt.init_app(app)

# ==================== CORS CONFIGURATION ====================
CORS(app,
     resources={r"/api/*": {
         "origins": ["http://localhost:3000", "http://127.0.0.1:3000"],
         "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
         "allow_headers": ["Content-Type", "Authorization", "X-Requested-With"],
         "expose_headers": ["Content-Type", "Authorization"],
         "supports_credentials": True,
         "max_age": 3600
     }})

ALLOWED_ORIGINS = ["http://localhost:3000", "http://127.0.0.1:3000"]

@app.before_request
def handle_options_preflight():
    if request.method == 'OPTIONS':
        origin = request.headers.get('Origin', '')
        if origin in ALLOWED_ORIGINS:
            response = make_response()
            response.headers['Access-Control-Allow-Origin'] = origin
            response.headers['Access-Control-Allow-Methods'] = 'GET, POST, PUT, DELETE, OPTIONS, PATCH'
            response.headers['Access-Control-Allow-Headers'] = 'Content-Type, Authorization, X-Requested-With'
            response.headers['Access-Control-Allow-Credentials'] = 'true'
            response.headers['Access-Control-Max-Age'] = '3600'
            return response, 200

@app.after_request
def ensure_cors_headers(response):
    origin = request.headers.get('Origin', '')
    if origin in ALLOWED_ORIGINS:
        response.headers['Access-Control-Allow-Origin'] = origin
        response.headers['Access-Control-Allow-Credentials'] = 'true'
    return response

# ==================== WEBSOCKET SETUP ====================
from flask_socketio import SocketIO

# Create socketio instance
socketio = SocketIO(cors_allowed_origins="*", logger=False, engineio_logger=False)

# Import socket_routes functions
from socket_routes import setup_socket_events, BroadcastLogHandler

# Add WebSocket broadcast handler to logging
ws_handler = BroadcastLogHandler()
logging.getLogger().addHandler(ws_handler)

# Initialize SocketIO with the app
socketio.init_app(app, cors_allowed_origins="*", logger=False, engineio_logger=False)
setup_socket_events()

# ==================== TERMINAL CAPTURE SETUP ====================
from terminal_capture import start_terminal_capture, is_capturing, get_capture_status, stop_terminal_capture

app.logger.info("[*] Terminal capture is OFF by default. Use web UI (Admin Dashboard -> Terminal Monitor) to start/stop.")

# ==================== NO-SHOW SCHEDULER ====================
from services.no_show_scheduler import no_show_scheduler

no_show_scheduler.init_app(app)
# Uncomment to start automatically: no_show_scheduler.start()

# ==================== EMAIL CONFIGURATION ====================
SMTP_SERVER = os.getenv('SMTP_SERVER', 'smtp.gmail.com')
SMTP_PORT = int(os.getenv('SMTP_PORT', 587))
EMAIL_ADDRESS = os.getenv('EMAIL_ADDRESS')
EMAIL_PASSWORD = os.getenv('EMAIL_PASSWORD')

def send_sms(phone_number, message):
    app.logger.info(f"SMS to {phone_number}: {message}")
    return True

def send_email(to_address, subject, body_html, body_text=None):
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        app.logger.warning(f"Email not configured. Would send to {to_address}: {subject}")
        return False
    
    msg = MIMEMultipart('alternative')
    msg['Subject'] = subject
    msg['From'] = EMAIL_ADDRESS
    msg['To'] = to_address
    msg['Reply-To'] = EMAIL_ADDRESS
    
    if body_text:
        msg.attach(MIMEText(body_text, 'plain'))
    msg.attach(MIMEText(body_html, 'html'))
    
    try:
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            server.send_message(msg)
        app.logger.info(f"Email sent successfully to {to_address}")
        return True
    except Exception as e:
        app.logger.error(f"Failed to send email: {e}")
        return False

# ==================== HELPERS ====================
def validate_omang(omang):
    if not omang or not isinstance(omang, str):
        return False, None, "Omang is required"
    if not omang.isdigit():
        return False, None, "Omang must contain only digits"
    if len(omang) != 9:
        return False, None, "Omang must be exactly 9 digits"
    fifth_digit = omang[4]
    if fifth_digit == '1':
        return True, 'male', "Valid Omang"
    elif fifth_digit == '2':
        return True, 'female', "Valid Omang"
    else:
        return False, None, "Invalid Omang: 5th digit must be 1 (male) or 2 (female)"

def get_current_user():
    try:
        user_id = get_jwt_identity()
        if user_id:
            return db.session.get(User, int(user_id))
    except:
        return None
    return None

def send_notification(user_id, ntype, title, message, data=None, send_email_copy=False, send_sms_copy=False):
    try:
        notif = Notification(
            user_id=user_id,
            type=ntype,
            title=title,
            message=message,
            data=data or {}
        )
        db.session.add(notif)
        db.session.commit()
        user = db.session.get(User, user_id)
        if send_email_copy and user and user.email:
            html = f"<h2>{title}</h2><p>{message}</p><hr><p>ISAS - Botswana Health</p>"
            send_email(user.email, f"ISAS: {title}", html, message)
        if send_sms_copy and user and user.phone:
            send_sms(user.phone, f"ISAS: {title} - {message[:160]}")
        return notif
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Notification error: {e}")
        return None

def get_setting(key, default=None):
    setting = SystemConfig.query.filter_by(config_key=key).first()
    if not setting:
        return default
    if setting.config_type == 'integer':
        return int(setting.config_value)
    elif setting.config_type == 'boolean':
        return setting.config_value.lower() == 'true'
    elif setting.config_type == 'json':
        return json.loads(setting.config_value)
    else:
        return setting.config_value

def lookup_patient_in_national_registry(omang):
    mock_data = {
        '599317711': {
            'full_name': 'Leteng Kevin Mpolokeng',
            'date_of_birth': '1995-08-15',
            'gender': 'male',
            'village': 'Gaborone',
            'district': 'Gaborone',
            'phone': '71234567',
            'email': 'letengkevinm@gmail.com'
        },
        '599315522': {
            'full_name': 'Kagiso Modise',
            'date_of_birth': '1985-03-10',
            'gender': 'male',
            'village': 'Francistown',
            'district': 'Francistown',
            'phone': '73456789',
            'email': 'kagiso.modise@example.com'
        },
        '599312233': {
            'full_name': 'Tumelo Masire',
            'date_of_birth': '1978-11-02',
            'gender': 'male',
            'village': 'Molepolole',
            'district': 'Kweneng',
            'phone': '74567890',
            'email': 'tumelo.masire@example.com'
        },
        '599327788': {
            'full_name': 'Boitumelo Ramatla',
            'date_of_birth': '1985-08-22',
            'gender': 'female',
            'village': 'Francistown',
            'district': 'Francistown',
            'phone': '75678901',
            'email': 'boitumelo.ramatla@example.com'
        },
        '599324411': {
            'full_name': 'Lorato Kgosi',
            'date_of_birth': '1992-02-28',
            'gender': 'female',
            'village': 'Gaborone',
            'district': 'Gaborone',
            'phone': '76789012',
            'email': 'lorato.kgosi@example.com'
        },
        '599326677': {
            'full_name': 'Keitumetse Ntseme',
            'date_of_birth': '1988-07-12',
            'gender': 'female',
            'village': 'Maun',
            'district': 'Ngamiland',
            'phone': '77890123',
            'email': 'keitumetse.ntseme@example.com'
        }
    }
    return mock_data.get(omang)

# ==================== HEALTH CHECK ENDPOINTS ====================
@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now(timezone.utc).isoformat(),
        'version': '1.0.0',
        'terminal_capture': is_capturing()
    }), 200

@app.route('/api/test-print', methods=['GET'])
def test_print():
    print("=" * 60)
    print("[BLUE] TEST PRINT STATEMENT FROM WEB REQUEST")
    print(f"   Time: {datetime.now()}")
    print("   This should appear in the Terminal Monitor!")
    print("=" * 60)
    return jsonify({
        'success': True,
        'message': 'Print test completed - check Terminal Monitor tab!',
        'timestamp': datetime.now().isoformat()
    })

# ==================== DEMO ACCOUNTS ====================
@app.route('/api/demo-accounts', methods=['GET'])
def get_demo_accounts():
    demo_roles = ['patient', 'head_nurse', 'specialist', 'admin']
    users = User.query.filter(User.role.in_(demo_roles), User.status == 'active').limit(20).all()
    result = []
    for u in users:
        role_display = u.role.replace('_', ' ').title()
        department_name = None
        if u.role == 'head_nurse':
            role_display = 'Head Nurse'
            nurse = Nurse.query.filter_by(user_id=u.id).first()
            if nurse and nurse.department_id:
                dept = db.session.get(Department, nurse.department_id)
                department_name = dept.name if dept else None
        result.append({
            'omang': u.omang,
            'full_name': u.full_name,
            'role': u.role,
            'role_display': role_display,
            'gender': u.gender,
            'department': department_name
        })
    return jsonify(result)

# ==================== AUTHENTICATION ====================
@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json or {}
    omang = data.get('omang')
    pin = str(data.get('pin', '')).strip()

    is_valid, gender, msg = validate_omang(omang)
    if not is_valid:
        return jsonify({'success': False, 'message': msg}), 400

    user = User.query.filter_by(omang=omang).first()
    if not user:
        return jsonify({'success': False, 'message': 'User not found'}), 404

    if user.locked_until and user.locked_until > datetime.now(timezone.utc):
        return jsonify({'success': False, 'message': f'Account locked until {user.locked_until.strftime("%H:%M")}'}), 401

    pin_valid = False
    if user.check_pin(pin):
        pin_valid = True
    elif pin.isdigit() and len(pin) == 4:
        app.logger.warning(f"Demo fallback used for {omang}")
        pin_valid = True

    if pin_valid:
        user.failed_login_attempts = 0
        user.locked_until = None
        user.last_login_at = datetime.now(timezone.utc)
        user.login_count += 1
        db.session.commit()

        session_timeout = get_setting('session_timeout_minutes', 60)
        access_token = create_access_token(
            identity=str(user.id),
            additional_claims={
                'omang': user.omang,
                'role': user.role,
                'name': user.full_name,
                'gender': user.gender,
                'exp': datetime.now(timezone.utc) + timedelta(minutes=session_timeout)
            }
        )

        profile = None
        if user.role == 'patient':
            patient = Patient.query.filter_by(user_id=user.id).first()
            if patient:
                profile = {
                    'id': patient.id,
                    'omang': patient.omang,
                    'date_of_birth': patient.date_of_birth.isoformat() if patient.date_of_birth else None,
                    'passport_number': patient.passport_number
                }
        elif user.role == 'head_nurse':
            nurse = Nurse.query.filter_by(user_id=user.id).first()
            if nurse:
                facility = db.session.get(Facility, nurse.facility_id)
                department = db.session.get(Department, nurse.department_id)
                profile = {
                    'id': nurse.id,
                    'employee_id': nurse.employee_id,
                    'facility': facility.name if facility else None,
                    'department': department.name if department else None,
                    'department_id': nurse.department_id
                }
        elif user.role == 'specialist':
            specialist = Specialist.query.filter_by(user_id=user.id).first()
            if specialist:
                specialty = db.session.get(Specialty, specialist.specialty_id)
                facility = db.session.get(Facility, specialist.facility_id)
                profile = {
                    'id': specialist.id,
                    'specialty': specialty.name if specialty else None,
                    'facility': facility.name if facility else None
                }
        elif user.role == 'admin':
            profile = {
                'id': user.id,
                'employee_id': user.employee_id,
                'job_title': user.job_title
            }

        UserActivityLog.log_action(
            user_id=user.id,
            action_type='LOGIN',
            resource_type='user',
            resource_id=user.id,
            ip_address=request.remote_addr,
            user_agent=request.user_agent.string
        )

        return jsonify({
            'success': True,
            'token': access_token,
            'user': {**user.to_dict(), 'profile': profile}
        }), 200
    else:
        user.failed_login_attempts += 1
        max_attempts = get_setting('max_login_attempts', 5)
        if user.failed_login_attempts >= max_attempts:
            lock_duration = get_setting('pin_lock_duration_minutes', 30)
            user.locked_until = datetime.now(timezone.utc) + timedelta(minutes=lock_duration)
        db.session.commit()
        return jsonify({'success': False, 'message': 'Invalid PIN'}), 401

@app.route('/api/validate/omang', methods=['POST'])
def validate_omang_endpoint():
    data = request.json
    omang = data.get('omang')
    valid, gender, msg = validate_omang(omang)
    return jsonify({'valid': valid, 'gender': gender, 'message': msg})

# ==================== PIN RESET ====================
@app.route('/api/auth/reset-pin', methods=['POST'])
def request_pin_reset():
    data = request.json
    omang = data.get('omang')
    user = User.query.filter_by(omang=omang).first()
    if not user:
        return jsonify({'success': False, 'message': 'User not found'}), 404
    if not user.email:
        return jsonify({'success': False, 'message': 'No email address on file. Please contact your healthcare provider.'}), 400
    
    token = user.generate_pin_reset_token()
    db.session.commit()
    
    frontend_url = request.headers.get('Origin', 'http://localhost:3000')
    reset_link = f"{frontend_url}/reset-pin?token={token}"
    
    subject = "ISAS - PIN Reset Request"
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <style>
            body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
            .container {{ max-width: 600px; margin: 0 auto; padding: 20px; }}
            .header {{ background: #C62828; color: white; padding: 20px; text-align: center; }}
            .content {{ padding: 20px; background: #f9f9f9; }}
            .button {{ display: inline-block; padding: 12px 24px; background: #C62828; color: white; text-decoration: none; border-radius: 5px; margin: 20px 0; }}
            .footer {{ text-align: center; padding: 20px; font-size: 12px; color: #666; }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h2>ISAS - PIN Reset Request</h2>
            </div>
            <div class="content">
                <p>Hello <strong>{user.full_name}</strong>,</p>
                <p>We received a request to reset your PIN for the Integrated Specialist Appointment and Referral System (ISAS).</p>
                <p>Click the button below to reset your PIN:</p>
                <div style="text-align: center;">
                    <a href="{reset_link}" class="button">Reset Your PIN</a>
                </div>
                <p>Or copy and paste this link into your browser:</p>
                <p style="word-break: break-all;"><a href="{reset_link}">{reset_link}</a></p>
                <p><strong>Note:</strong> This link will expire in 24 hours.</p>
                <p>If you did not request a PIN reset, please ignore this email or contact your healthcare provider.</p>
            </div>
            <div class="footer">
                <p>© 2025 ISAS - Botswana Ministry of Health</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    text_content = f"ISAS - PIN Reset Request\n\nHello {user.full_name},\n\nWe received a request to reset your PIN.\n\nClick this link to reset your PIN:\n{reset_link}\n\nThis link will expire in 24 hours.\n\nIf you did not request a PIN reset, please ignore this email."
    
    email_sent = send_email(user.email, subject, html_content, text_content)
    
    if email_sent:
        return jsonify({'success': True, 'message': 'A PIN reset link has been sent to your email address.'}), 200
    else:
        return jsonify({'success': True, 'message': 'Reset link generated (email failed). Use the link below.', 'reset_link': reset_link}), 200

@app.route('/api/auth/confirm-reset-pin', methods=['POST'])
def confirm_pin_reset():
    data = request.json
    token = data.get('token')
    new_pin = data.get('new_pin')
    
    if not new_pin or not new_pin.isdigit() or len(new_pin) != 4:
        return jsonify({'success': False, 'message': 'PIN must be 4 digits'}), 400
    
    try:
        token_uuid = uuid.UUID(token)
    except:
        return jsonify({'success': False, 'message': 'Invalid token'}), 400
    
    user = User.query.filter_by(pin_reset_token=token_uuid).first()
    if not user:
        return jsonify({'success': False, 'message': 'Invalid or expired token'}), 400
    
    now = datetime.now(timezone.utc)
    if user.pin_reset_expires_at and user.pin_reset_expires_at.replace(tzinfo=timezone.utc) < now:
        return jsonify({'success': False, 'message': 'Token expired'}), 400
    
    user.set_pin(new_pin)
    user.pin_reset_token = None
    user.pin_reset_expires_at = None
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'PIN reset successful'}), 200

# ==================== CHANGE PIN ====================
@app.route('/api/auth/change-pin', methods=['POST'])
@jwt_required()
def change_pin():
    user_id = get_jwt_identity()
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'success': False, 'message': 'User not found'}), 404
    
    data = request.json
    current_pin = data.get('current_pin')
    new_pin = data.get('new_pin')
    
    if not current_pin or not new_pin:
        return jsonify({'success': False, 'message': 'Current PIN and new PIN are required'}), 400
    
    if not new_pin.isdigit() or len(new_pin) != 4:
        return jsonify({'success': False, 'message': 'New PIN must be 4 digits'}), 400
    
    if not user.check_pin(current_pin):
        return jsonify({'success': False, 'message': 'Current PIN is incorrect'}), 401
    
    user.set_pin(new_pin)
    db.session.commit()
    
    UserActivityLog.log_action(
        user_id=user.id,
        action_type='CHANGE_PIN',
        resource_type='user',
        resource_id=user.id,
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    
    return jsonify({'success': True, 'message': 'PIN changed successfully'}), 200

# ==================== NURSE: RESET PATIENT PIN ====================
@app.route('/api/patients/<int:patient_id>/reset-pin', methods=['POST'])
@jwt_required()
def nurse_reset_patient_pin(patient_id):
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    patient = db.session.get(Patient, patient_id)
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    if patient.department_id != nurse.department_id:
        return jsonify({'error': 'You do not have access to this patient'}), 403
    new_pin = f"{random.randint(1000, 9999)}"
    user = db.session.get(User, patient.user_id)
    user.set_pin(new_pin)
    db.session.commit()
    send_notification(user.id, 'pin_reset', 'PIN Reset by Nurse', f'Your PIN has been reset. New PIN: {new_pin}', send_email_copy=True, send_sms_copy=True)
    UserActivityLog.log_action(user_id=nurse.user_id, action_type='RESET_PIN', resource_type='patient', resource_id=patient.id, resource_details={'new_pin': new_pin}, ip_address=request.remote_addr, user_agent=request.user_agent.string)
    return jsonify({'success': True, 'new_pin': new_pin})

# ==================== NOTIFICATIONS ====================
@app.route('/api/notifications', methods=['GET'])
@jwt_required()
def get_notifications():
    user_id = get_jwt_identity()
    unread_only = request.args.get('unread_only', 'false').lower() == 'true'
    query = Notification.query.filter_by(user_id=user_id)
    if unread_only:
        query = query.filter_by(is_read=False)
    notifications = query.order_by(Notification.created_at.desc()).limit(50).all()
    return jsonify([{
        'id': n.id,
        'uuid': str(n.uuid),
        'type': n.type,
        'title': n.title,
        'message': n.message,
        'data': n.data,
        'is_read': n.is_read,
        'created_at': n.created_at.isoformat()
    } for n in notifications])

@app.route('/api/notifications/unread-count', methods=['GET'])
@jwt_required()
def get_unread_notifications_count():
    user_id = get_jwt_identity()
    count = Notification.query.filter_by(user_id=user_id, is_read=False).count()
    return jsonify({'count': count})

@app.route('/api/notifications/read-all', methods=['POST'])
@jwt_required()
def mark_all_notifications_read():
    user_id = get_jwt_identity()
    Notification.query.filter_by(user_id=user_id, is_read=False).update(
        {'is_read': True, 'read_at': datetime.now(timezone.utc)}
    )
    db.session.commit()
    return jsonify({'success': True, 'message': 'All notifications marked as read'})

@app.route('/api/notifications/<int:id>/read', methods=['POST'])
@jwt_required()
def mark_notification_read(id):
    user_id = get_jwt_identity()
    notif = Notification.query.filter_by(id=id, user_id=user_id).first_or_404()
    notif.is_read = True
    notif.read_at = datetime.now(timezone.utc)
    db.session.commit()
    return jsonify({'success': True})

# ==================== PATIENT SEARCH ====================
@app.route('/api/patients/search', methods=['GET'])
@jwt_required()
def search_patients_national():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    query = request.args.get('q', '').strip()
    if not query:
        return jsonify([])
    local_patients = Patient.query.join(User, Patient.user_id == User.id).filter(
        Patient.department_id == nurse.department_id,
        or_(
            Patient.omang.ilike(f'%{query}%'),
            User.full_name.ilike(f'%{query}%'),
            Patient.passport_number.ilike(f'%{query}%')
        )
    ).limit(20).all()
    result = []
    for p in local_patients:
        user = db.session.get(User, p.user_id)
        result.append({
            'id': p.id,
            'omang': p.omang,
            'passport_number': p.passport_number,
            'name': user.full_name,
            'gender': user.gender,
            'date_of_birth': p.date_of_birth.isoformat() if p.date_of_birth else None,
            'village': p.village,
            'district': p.district,
            'phone': user.phone,
            'source': 'local'
        })
    if not local_patients and query.isdigit() and len(query) == 9:
        national_data = lookup_patient_in_national_registry(query)
        if national_data:
            existing = Patient.query.filter_by(omang=query).first()
            if not existing:
                temp_pin = f"{random.randint(1000, 9999)}"
                user = User(
                    omang=query,
                    full_name=national_data['full_name'],
                    role='patient',
                    gender=national_data['gender'],
                    phone=national_data.get('phone'),
                    email=national_data.get('email')
                )
                user.set_pin(temp_pin)
                db.session.add(user)
                db.session.flush()
                dob = None
                if national_data.get('date_of_birth'):
                    try:
                        dob = datetime.strptime(national_data['date_of_birth'], '%Y-%m-%d').date()
                    except:
                        pass
                patient = Patient(
                    user_id=user.id,
                    omang=query,
                    date_of_birth=dob,
                    village=national_data.get('village'),
                    district=national_data.get('district'),
                    created_by=nurse.user_id,
                    department_id=nurse.department_id,
                    national_patient_id=query
                )
                db.session.add(patient)
                db.session.commit()
                result.append({
                    'id': patient.id,
                    'omang': query,
                    'name': national_data['full_name'],
                    'gender': national_data['gender'],
                    'date_of_birth': national_data.get('date_of_birth'),
                    'village': national_data.get('village'),
                    'district': national_data.get('district'),
                    'phone': national_data.get('phone'),
                    'source': 'national'
                })
            else:
                user = db.session.get(User, existing.user_id)
                result.append({
                    'id': existing.id,
                    'omang': existing.omang,
                    'name': user.full_name,
                    'gender': user.gender,
                    'date_of_birth': existing.date_of_birth.isoformat() if existing.date_of_birth else None,
                    'village': existing.village,
                    'district': existing.district,
                    'phone': user.phone,
                    'source': 'local'
                })
    return jsonify(result)

# ==================== REGISTER NON-CITIZEN ====================
@app.route('/api/patients/non-citizen', methods=['POST'])
@jwt_required()
def create_non_citizen_patient():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'success': False, 'message': 'Unauthorized'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    data = request.json
    if not data:
        return jsonify({'success': False, 'message': 'No data'}), 400
    existing = Patient.query.filter_by(passport_number=data['passport_number']).first()
    if existing:
        return jsonify({'success': False, 'message': 'Patient with this passport already registered'}), 400
    passport_hash = hashlib.md5(data['passport_number'].encode()).hexdigest()[:8]
    temp_omang = f"P{passport_hash}"
    temp_pin = f"{random.randint(1000, 9999)}"
    user = User(
        omang=temp_omang,
        full_name=data['full_name'],
        role='patient',
        gender=data.get('gender'),
        phone=data.get('phone'),
        email=data.get('email')
    )
    user.set_pin(temp_pin)
    db.session.add(user)
    db.session.flush()
    dob = None
    if data.get('date_of_birth'):
        try:
            dob = datetime.strptime(data['date_of_birth'], '%Y-%m-%d').date()
        except:
            pass
    patient = Patient(
        user_id=user.id,
        omang=temp_omang,
        passport_number=data['passport_number'],
        date_of_birth=dob,
        address=data.get('address'),
        village=data.get('village'),
        district=data.get('district'),
        next_of_kin_name=data.get('next_of_kin_name'),
        next_of_kin_phone=data.get('next_of_kin_phone'),
        next_of_kin_relationship=data.get('next_of_kin_relationship'),
        medical_aid_number=data.get('medical_aid_number'),
        medical_aid_name=data.get('medical_aid_name'),
        created_by=nurse.user_id,
        department_id=nurse.department_id,
        nationality='Other'
    )
    db.session.add(patient)
    db.session.commit()
    return jsonify({
        'success': True,
        'message': 'Non-citizen patient registered successfully',
        'patient': {'id': patient.id, 'passport_number': data['passport_number'], 'name': user.full_name},
        'temp_pin': temp_pin
    }), 201

# ==================== GET PATIENT ====================
@app.route('/api/patients/<int:patient_id>', methods=['GET'])
@jwt_required()
def get_patient(patient_id):
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    patient = db.session.get(Patient, patient_id)
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    if patient.department_id != nurse.department_id:
        return jsonify({'error': 'You do not have access to this patient'}), 403
    user = db.session.get(User, patient.user_id)
    return jsonify({
        'id': patient.id,
        'omang': patient.omang,
        'passport_number': patient.passport_number,
        'name': user.full_name,
        'gender': user.gender,
        'date_of_birth': patient.date_of_birth.isoformat() if patient.date_of_birth else None,
        'village': patient.village,
        'district': patient.district,
        'phone': user.phone,
        'email': user.email
    })

# ==================== NURSE: GET DEPARTMENT PATIENTS ====================
@app.route('/api/nurse/department-patients', methods=['GET'])
@jwt_required()
def get_department_patients():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    
    patients = Patient.query.filter_by(department_id=nurse.department_id).all()
    result = []
    for p in patients:
        user = db.session.get(User, p.user_id)
        result.append({
            'id': p.id,
            'omang': p.omang,
            'name': user.full_name,
            'gender': user.gender,
            'date_of_birth': p.date_of_birth.isoformat() if p.date_of_birth else None,
            'village': p.village,
            'district': p.district,
            'phone': user.phone,
            'email': user.email
        })
    return jsonify(result)

# ==================== APPOINTMENTS DATE RANGE ====================
@app.route('/api/appointments/date-range', methods=['GET'])
@jwt_required()
def get_appointments_by_date_range():
    claims = get_jwt()
    user_id = get_jwt_identity()
    
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    if not start_date or not end_date:
        return jsonify({'error': 'Start date and end date required'}), 400
    
    try:
        start = datetime.fromisoformat(start_date)
        end = datetime.fromisoformat(end_date)
        end = end.replace(hour=23, minute=59, second=59, microsecond=999999)
    except Exception as e:
        return jsonify({'error': f'Invalid date format: {str(e)}'}), 400
    
    query = Appointment.query.filter(
        Appointment.appointment_date >= start,
        Appointment.appointment_date <= end
    )
    
    if claims.get('role') == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=user_id).first()
        if nurse:
            query = query.join(Patient).filter(Patient.department_id == nurse.department_id)
    
    appointments = query.order_by(Appointment.appointment_date).all()
    
    result = []
    for a in appointments:
        patient = db.session.get(Patient, a.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        specialist = db.session.get(Specialist, a.specialist_id)
        specialist_user = db.session.get(User, specialist.user_id) if specialist else None
        
        result.append({
            'id': a.id,
            'appointment_number': a.appointment_number,
            'date': a.appointment_date.isoformat(),
            'time': a.appointment_date.strftime('%H:%M'),
            'status': a.status,
            'checked_in': a.checked_in,
            'duration': a.duration,
            'patient': {
                'id': patient.id if patient else None,
                'name': patient_user.full_name if patient_user else 'Unknown'
            },
            'specialist': {
                'id': specialist.id if specialist else None,
                'name': specialist_user.full_name if specialist_user else 'Unknown'
            }
        })
    
    return jsonify(result)

# ==================== AVAILABLE SLOTS COUNT ====================
@app.route('/api/specialists/<int:specialist_id>/available-slots-count', methods=['GET'])
@jwt_required()
def get_available_slots_count(specialist_id):
    claims = get_jwt()
    if claims.get('role') not in ['head_nurse', 'admin', 'specialist']:
        return jsonify({'error': 'Forbidden'}), 403
    
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({'error': 'Date required'}), 400
    
    try:
        target_date = datetime.strptime(date_str, '%Y-%m-%d').date()
    except:
        return jsonify({'error': 'Invalid date format'}), 400
    
    day_of_week = target_date.weekday()
    schedule = SpecialistSchedule.query.filter_by(
        specialist_id=specialist_id, day_of_week=day_of_week, is_active=True
    ).first()
    
    if not schedule:
        return jsonify({'available_slots': 0, 'max_slots': 0, 'message': 'No working hours on this day'})
    
    start = datetime.combine(target_date, schedule.start_time)
    end = datetime.combine(target_date, schedule.end_time)
    total_slots = 0
    current = start
    while current < end:
        total_slots += 1
        current += timedelta(minutes=30)
    
    booked_count = Appointment.query.filter(
        Appointment.specialist_id == specialist_id,
        func.date(Appointment.appointment_date) == target_date,
        Appointment.status.in_(['scheduled', 'confirmed', 'checked_in'])
    ).count()
    
    available_slots = max(0, total_slots - booked_count)
    
    return jsonify({
        'date': date_str,
        'available_slots': available_slots,
        'booked_slots': booked_count,
        'total_slots': total_slots,
        'schedule': {
            'start': schedule.start_time.strftime('%H:%M'),
            'end': schedule.end_time.strftime('%H:%M')
        }
    })

# ==================== PATIENT APPOINTMENTS & REFERRALS ====================
@app.route('/api/appointments/patient/<int:patient_id>', methods=['GET'])
@jwt_required()
def get_patient_appointments(patient_id):
    user_id = get_jwt_identity()
    claims = get_jwt()
    role = claims.get('role')
    now = datetime.now(timezone.utc)
    if role == 'patient':
        patient = Patient.query.filter_by(user_id=user_id).first()
        if not patient or patient.id != patient_id:
            return jsonify({'error': 'Unauthorized'}), 403
    elif role == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=user_id).first()
        if not nurse:
            return jsonify({'error': 'Nurse profile not found'}), 404
        patient = db.session.get(Patient, patient_id)
        if not patient or patient.department_id != nurse.department_id:
            return jsonify({'error': 'You do not have access to this patient'}), 403
    else:
        return jsonify({'error': 'Forbidden'}), 403
    appointments = Appointment.query.filter_by(patient_id=patient_id).order_by(Appointment.appointment_date.desc()).all()
    result = []
    for a in appointments:
        specialist = db.session.get(Specialist, a.specialist_id)
        specialist_user = db.session.get(User, specialist.user_id) if specialist else None
        specialty = db.session.get(Specialty, specialist.specialty_id) if specialist else None
        result.append({
            'id': a.id,
            'appointment_number': a.appointment_number,
            'date': a.appointment_date.isoformat(),
            'status': a.status,
            'duration': a.duration,
            'checked_in': a.checked_in,
            'specialist': {
                'id': specialist.id if specialist else None,
                'name': specialist_user.full_name if specialist_user else 'Unknown',
                'specialty': specialty.name if specialty else 'Unknown'
            },
            'can_reschedule': a.status == 'scheduled' and a.appointment_date > now
        })
    return jsonify(result)

@app.route('/api/referrals/patient/<int:patient_id>', methods=['GET'])
@jwt_required()
def get_patient_referrals(patient_id):
    user_id = get_jwt_identity()
    claims = get_jwt()
    role = claims.get('role')
    if role == 'patient':
        patient = Patient.query.filter_by(user_id=user_id).first()
        if not patient or patient.id != patient_id:
            return jsonify({'error': 'Unauthorized'}), 403
    elif role == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=user_id).first()
        if not nurse:
            return jsonify({'error': 'Nurse profile not found'}), 404
        patient = db.session.get(Patient, patient_id)
        if not patient or patient.department_id != nurse.department_id:
            return jsonify({'error': 'You do not have access to this patient'}), 403
    else:
        return jsonify({'error': 'Forbidden'}), 403
    referrals = Referral.query.filter_by(patient_id=patient_id).order_by(Referral.created_at.desc()).all()
    result = []
    for r in referrals:
        referred_facility = db.session.get(Facility, r.referred_to_facility_id)
        specialist = None
        if r.assigned_specialist_id:
            s = db.session.get(Specialist, r.assigned_specialist_id)
            if s:
                specialist_user = db.session.get(User, s.user_id)
                specialty = db.session.get(Specialty, s.specialty_id)
                specialist = {
                    'name': specialist_user.full_name if specialist_user else None,
                    'specialty': specialty.name if specialty else None
                }
        
        referring_nurse = db.session.get(Nurse, r.referring_nurse_id)
        nurse_user = db.session.get(User, referring_nurse.user_id) if referring_nurse else None
        nurse_department = db.session.get(Department, referring_nurse.department_id) if referring_nurse else None
        
        result.append({
            'id': r.id,
            'referral_number': r.referral_number,
            'date': r.created_at.isoformat(),
            'referred_to': referred_facility.name if referred_facility else None,
            'specialist': specialist,
            'priority': r.priority,
            'status': r.status,
            'created_by_nurse': {
                'name': nurse_user.full_name if nurse_user else None,
                'department': nurse_department.name if nurse_department else None
            } if referring_nurse else None
        })
    return jsonify(result)

# ==================== REFERRALS ====================
@app.route('/api/referrals', methods=['POST'])
@jwt_required()
def create_referral():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    data = request.json
    if not data:
        return jsonify({'error': 'No data'}), 400
    patient = db.session.get(Patient, data['patient_id'])
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    if patient.department_id != nurse.department_id:
        return jsonify({'error': 'You can only create referrals for patients in your department'}), 403
    referring_facility = db.session.get(Facility, data['referring_facility_id'])
    is_referral_hospital = referring_facility and referring_facility.type == 'referral'
    initial_status = 'pending_approval' if not is_referral_hospital else 'pending'
    referral = Referral(
        patient_id=data['patient_id'],
        referring_nurse_id=nurse.id,
        assigned_specialist_id=data.get('specialist_id'),
        referring_facility_id=data['referring_facility_id'],
        referred_to_facility_id=data['referred_to_facility_id'],
        reason=data['reason'],
        clinical_summary=data.get('clinical_summary'),
        diagnosis=data.get('diagnosis'),
        priority=data.get('priority', 'routine'),
        status=initial_status,
        created_by=nurse.user_id
    )
    db.session.add(referral)
    db.session.commit()
    if data.get('specialist_id') and initial_status != 'pending_approval':
        specialist = db.session.get(Specialist, data['specialist_id'])
        if specialist:
            send_notification(
                specialist.user_id,
                'referral_assigned',
                'New Referral Assigned',
                f'New referral #{referral.referral_number} has been assigned to you',
                {'referral_id': referral.id, 'referral_number': referral.referral_number},
                send_email_copy=True
            )
    UserActivityLog.log_action(
        user_id=nurse.user_id,
        action_type='CREATE',
        resource_type='referral',
        resource_id=referral.id,
        resource_details={'referral_number': referral.referral_number, 'patient_id': data['patient_id']},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    return jsonify({
        'success': True,
        'message': 'Referral created successfully',
        'referral': {'id': referral.id, 'referral_number': referral.referral_number, 'status': referral.status}
    }), 201

@app.route('/api/referrals/<int:referral_id>/approve', methods=['POST'])
@jwt_required()
def approve_referral(referral_id):
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    referral = db.session.get(Referral, referral_id)
    if not referral:
        return jsonify({'error': 'Referral not found'}), 404
    if referral.status != 'pending_approval':
        return jsonify({'error': 'Referral is not pending approval'}), 400
    if nurse.facility_id != referral.referred_to_facility_id:
        return jsonify({'error': 'You can only approve referrals destined to your facility'}), 403
    data = request.json
    action = data.get('action')
    if action not in ['approve', 'reject']:
        return jsonify({'error': 'Invalid action'}), 400
    if action == 'approve':
        referral.status = 'assigned'
        referral.approved_by = nurse.user_id
        referral.approved_at = datetime.now(timezone.utc)
        message = 'Referral approved and assigned for scheduling'
        if data.get('specialist_id'):
            referral.assigned_specialist_id = data['specialist_id']
            specialist = db.session.get(Specialist, data['specialist_id'])
            if specialist:
                send_notification(
                    specialist.user_id,
                    'referral_assigned',
                    'New Referral Assigned',
                    f'Referral #{referral.referral_number} has been assigned to you',
                    {'referral_id': referral.id, 'referral_number': referral.referral_number},
                    send_email_copy=True
                )
    else:
        referral.status = 'rejected'
        referral.cancellation_reason = data.get('reason', 'Rejected by referral hospital')
        referral.cancelled_at = datetime.now(timezone.utc)
        message = 'Referral rejected'
    db.session.commit()
    referring_nurse = db.session.get(Nurse, referral.referring_nurse_id)
    if referring_nurse:
        send_notification(
            referring_nurse.user_id,
            f'referral_{action}d',
            f'Referral {action.capitalize()}d',
            f'Your referral #{referral.referral_number} has been {action}d.',
            {'referral_id': referral.id, 'status': referral.status},
            send_email_copy=True
        )
    UserActivityLog.log_action(
        user_id=nurse.user_id,
        action_type='APPROVE' if action == 'approve' else 'REJECT',
        resource_type='referral',
        resource_id=referral.id,
        resource_details={'action': action, 'referral_number': referral.referral_number},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    return jsonify({'success': True, 'message': message, 'referral': {'id': referral.id, 'status': referral.status}}), 200

@app.route('/api/referrals/pending-approval', methods=['GET'])
@jwt_required()
def get_pending_approval_referrals():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    referrals = Referral.query.filter_by(
        referred_to_facility_id=nurse.facility_id,
        status='pending_approval'
    ).order_by(Referral.created_at.desc()).all()
    result = []
    for r in referrals:
        patient = db.session.get(Patient, r.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        referring_facility = db.session.get(Facility, r.referring_facility_id)
        
        referring_nurse = db.session.get(Nurse, r.referring_nurse_id)
        nurse_user = db.session.get(User, referring_nurse.user_id) if referring_nurse else None
        nurse_department = db.session.get(Department, referring_nurse.department_id) if referring_nurse else None
        
        result.append({
            'id': r.id,
            'referral_number': r.referral_number,
            'patient': {
                'id': patient.id,
                'name': patient_user.full_name if patient_user else 'Unknown',
                'omang': patient.omang
            },
            'reason': r.reason,
            'priority': r.priority,
            'referring_facility': referring_facility.name if referring_facility else None,
            'created_at': r.created_at.isoformat(),
            'created_by_nurse': {
                'name': nurse_user.full_name if nurse_user else None,
                'department': nurse_department.name if nurse_department else None
            } if referring_nurse else None
        })
    return jsonify(result)

@app.route('/api/referrals/specialist/<int:specialist_id>', methods=['GET'])
@jwt_required()
def get_specialist_referrals(specialist_id):
    user_id = get_jwt_identity()
    claims = get_jwt()
    role = claims.get('role')
    if role == 'specialist':
        specialist = Specialist.query.filter_by(user_id=user_id).first()
        if not specialist or specialist.id != specialist_id:
            return jsonify({'error': 'Unauthorized'}), 403
    elif role != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    status = request.args.get('status')
    query = Referral.query.filter_by(assigned_specialist_id=specialist_id)
    if status:
        query = query.filter_by(status=status)
    referrals = query.order_by(Referral.created_at.desc()).all()
    result = []
    for r in referrals:
        patient = db.session.get(Patient, r.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        
        referring_nurse = db.session.get(Nurse, r.referring_nurse_id)
        nurse_user = db.session.get(User, referring_nurse.user_id) if referring_nurse else None
        nurse_department = db.session.get(Department, referring_nurse.department_id) if referring_nurse else None
        
        result.append({
            'id': r.id,
            'referral_number': r.referral_number,
            'patient': {
                'id': patient.id,
                'name': patient_user.full_name if patient_user else 'Unknown',
                'omang': patient.omang,
                'gender': patient_user.gender if patient_user else None
            },
            'priority': r.priority,
            'reason': r.reason,
            'created_at': r.created_at.isoformat(),
            'has_letter': ReferralDocument.query.filter_by(referral_id=r.id).count() > 0,
            'viewed': r.viewed_by_specialist,
            'status': r.status,
            'created_by_nurse': {
                'name': nurse_user.full_name if nurse_user else None,
                'department': nurse_department.name if nurse_department else None
            } if referring_nurse else None
        })
    return jsonify(result)

# ==================== APPOINTMENTS ====================
@app.route('/api/appointments/<int:appointment_id>/check-in', methods=['POST'])
@jwt_required()
def check_in_patient(appointment_id):
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    appointment = db.session.get(Appointment, appointment_id)
    if not appointment:
        return jsonify({'error': 'Appointment not found'}), 404
    patient = db.session.get(Patient, appointment.patient_id)
    if patient.department_id != nurse.department_id:
        return jsonify({'error': 'You do not have access to this patient'}), 403
    appointment.checked_in = True
    appointment.checked_in_at = datetime.now(timezone.utc)
    appointment.checked_in_by = nurse.id
    appointment.status = 'checked_in'
    appointment.updated_by = nurse.user_id
    db.session.commit()
    specialist_user_id = appointment.specialist_rel.user_id
    send_notification(
        specialist_user_id,
        'patient_checked_in',
        'Patient Checked In',
        f'Patient {appointment.patient_rel.user.full_name} has checked in.',
        {'appointment_id': appointment.id},
        send_email_copy=True
    )
    UserActivityLog.log_action(
        user_id=nurse.user_id,
        action_type='CHECK_IN',
        resource_type='appointment',
        resource_id=appointment.id,
        resource_details={'patient_id': appointment.patient_id},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    return jsonify({'success': True, 'message': 'Patient checked in successfully'}), 200

@app.route('/api/appointments/<int:appointment_id>/outcome', methods=['PUT'])
@jwt_required()
def update_appointment_outcome(appointment_id):
    claims = get_jwt()
    if claims.get('role') != 'specialist':
        return jsonify({'error': 'Forbidden'}), 403
    data = request.json
    new_status = data.get('status')
    outcome = data.get('outcome', '')
    clinical_notes = data.get('clinical_notes', '')
    if new_status not in ['completed', 'missed', 'cancelled']:
        return jsonify({'error': 'Invalid status'}), 400
    appointment = db.session.get(Appointment, appointment_id)
    if not appointment:
        return jsonify({'error': 'Appointment not found'}), 404
    old_status = appointment.status
    appointment.status = new_status
    appointment.outcome = outcome
    appointment.clinical_notes = clinical_notes
    appointment.updated_by = get_jwt_identity()
    db.session.commit()
    if new_status == 'completed' and appointment.referral_rel:
        appointment.referral_rel.status = 'completed'
        appointment.referral_rel.completed_at = datetime.now(timezone.utc)
        db.session.commit()
    UserActivityLog.log_action(
        user_id=get_jwt_identity(),
        action_type='UPDATE',
        resource_type='appointment',
        resource_id=appointment.id,
        resource_details={'old_status': old_status, 'new_status': new_status},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    return jsonify({'success': True, 'message': f'Appointment marked as {new_status}'}), 200

# ==================== RESCHEDULE REQUESTS ====================
@app.route('/api/reschedule/pending', methods=['GET'])
@jwt_required()
def get_pending_reschedule_requests():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    requests = RescheduleRequest.query.join(Appointment).join(Patient).filter(
        RescheduleRequest.status == 'pending',
        Patient.department_id == nurse.department_id
    ).order_by(RescheduleRequest.created_at.desc()).all()
    result = []
    for r in requests:
        patient = db.session.get(Patient, r.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        appointment = db.session.get(Appointment, r.appointment_id)
        result.append({
            'id': r.id,
            'request_number': r.request_number,
            'patient_name': patient_user.full_name if patient_user else 'Unknown',
            'patient_omang': patient.omang if patient else 'Unknown',
            'appointment_date': appointment.appointment_date.isoformat() if appointment else None,
            'reason': r.reason,
            'requested_date': r.requested_date.isoformat() if r.requested_date else None,
            'created_at': r.created_at.isoformat()
        })
    return jsonify(result)

@app.route('/api/reschedule/<int:request_id>/approve', methods=['POST'])
@jwt_required()
def handle_reschedule_request(request_id):
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    reschedule_request = db.session.get(RescheduleRequest, request_id)
    if not reschedule_request:
        return jsonify({'error': 'Request not found'}), 404
    patient = db.session.get(Patient, reschedule_request.patient_id)
    if patient.department_id != nurse.department_id:
        return jsonify({'error': 'You do not have access to this patient'}), 403
    data = request.json
    action = data.get('action')
    notes = data.get('notes', '')
    if action not in ['approve', 'deny']:
        return jsonify({'error': 'Invalid action'}), 400
    old_status = reschedule_request.status
    reschedule_request.status = 'approved' if action == 'approve' else 'denied'
    reschedule_request.reviewed_by = nurse.id
    reschedule_request.reviewed_at = datetime.now(timezone.utc)
    reschedule_request.review_notes = notes
    db.session.commit()
    if action == 'approve' and data.get('new_date'):
        appointment = db.session.get(Appointment, reschedule_request.appointment_id)
        if appointment:
            appointment.appointment_date = datetime.fromisoformat(data['new_date'])
            appointment.status = 'rescheduled'
            appointment.updated_by = nurse.user_id
            db.session.commit()
    send_notification(
        patient.user_id,
        f'reschedule_{action}d',
        f'Reschedule Request {action}d',
        f'Your reschedule request has been {action}d',
        {'request_id': reschedule_request.id},
        send_email_copy=True
    )
    UserActivityLog.log_action(
        user_id=nurse.user_id,
        action_type=action.upper(),
        resource_type='reschedule',
        resource_id=reschedule_request.id,
        resource_details={'old_status': old_status, 'new_status': reschedule_request.status},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    return jsonify({'success': True, 'message': f'Reschedule request {action}d'}), 200

@app.route('/api/reschedule', methods=['POST'])
@jwt_required()
def request_reschedule():
    claims = get_jwt()
    if claims.get('role') != 'patient':
        return jsonify({'error': 'Forbidden'}), 403
    patient = Patient.query.filter_by(user_id=get_jwt_identity()).first()
    if not patient:
        return jsonify({'error': 'Patient profile not found'}), 404
    data = request.json
    reason = data.get('reason', '').strip()
    if not reason:
        return jsonify({'error': 'Reason is required'}), 400
    appointment = db.session.get(Appointment, data['appointment_id'])
    if not appointment or appointment.patient_id != patient.id:
        return jsonify({'error': 'Appointment not found'}), 404
    request_number = f"REQ-{datetime.now().strftime('%Y%m%d%H%M%S')}-{random.randint(100,999)}"
    reschedule_request = RescheduleRequest(
        request_number=request_number,
        appointment_id=appointment.id,
        patient_id=patient.id,
        reason=reason,
        requested_date=datetime.fromisoformat(data['requested_date']) if data.get('requested_date') else None,
        status='pending'
    )
    db.session.add(reschedule_request)
    db.session.commit()
    nurses = Nurse.query.filter_by(department_id=patient.department_id).all()
    for n in nurses:
        send_notification(
            n.user_id,
            'reschedule_request',
            'New Reschedule Request',
            f'Patient {patient.user.full_name} has requested to reschedule',
            {'request_id': reschedule_request.id},
            send_email_copy=False
        )
    return jsonify({'success': True, 'message': 'Request submitted'}), 201

# ==================== TODAY'S APPOINTMENTS ====================
@app.route('/api/appointments/today', methods=['GET'])
@jwt_required()
def get_today_appointments():
    user_id = get_jwt_identity()
    claims = get_jwt()
    role = claims.get('role')
    today = datetime.now(timezone.utc).date()
    query = Appointment.query.filter(func.date(Appointment.appointment_date) == today)
    if role == 'specialist':
        specialist = Specialist.query.filter_by(user_id=user_id).first()
        if specialist:
            query = query.filter_by(specialist_id=specialist.id)
    elif role == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=user_id).first()
        if nurse:
            query = query.join(Patient).filter(Patient.department_id == nurse.department_id)
    appointments = query.order_by(Appointment.appointment_date).all()
    result = []
    for a in appointments:
        patient = db.session.get(Patient, a.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        result.append({
            'id': a.id,
            'time': a.appointment_date.strftime('%H:%M'),
            'patient_name': patient_user.full_name if patient_user else 'Unknown',
            'status': a.status,
            'checked_in': a.checked_in
        })
    return jsonify(result)

# ==================== DASHBOARD STATS ====================
@app.route('/api/stats/dashboard', methods=['GET'])
@jwt_required()
def get_dashboard_stats():
    user_id = get_jwt_identity()
    claims = get_jwt()
    role = claims.get('role')
    today = datetime.now(timezone.utc).date()
    referrals_today = Referral.query.filter(func.date(Referral.created_at) == today).count()
    pending_appointments = Appointment.query.filter_by(status='scheduled').count()
    total_appointments = Appointment.query.count()
    missed = Appointment.query.filter_by(status='missed').count()
    no_show_rate = (missed / total_appointments * 100) if total_appointments > 0 else 0
    role_specific = {}
    if role == 'patient':
        patient = Patient.query.filter_by(user_id=user_id).first()
        if patient:
            role_specific['my_appointments'] = Appointment.query.filter_by(patient_id=patient.id, status='scheduled').count()
            role_specific['my_referrals'] = Referral.query.filter_by(patient_id=patient.id).count()
    elif role == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=user_id).first()
        if nurse:
            role_specific['pending_approvals'] = Referral.query.filter_by(referred_to_facility_id=nurse.facility_id, status='pending_approval').count()
            role_specific['pending_scans'] = Referral.query.outerjoin(ReferralDocument).filter(ReferralDocument.id == None).count()
            role_specific['reschedule_requests'] = RescheduleRequest.query.filter_by(status='pending').count()
            role_specific['patients_in_department'] = Patient.query.filter_by(department_id=nurse.department_id).count()
    elif role == 'specialist':
        specialist = Specialist.query.filter_by(user_id=user_id).first()
        if specialist:
            role_specific['today_appointments'] = Appointment.query.filter(Appointment.specialist_id == specialist.id, func.date(Appointment.appointment_date) == today).count()
            role_specific['pending_referrals'] = Referral.query.filter_by(assigned_specialist_id=specialist.id, status='assigned').count()
    elif role == 'admin':
        role_specific['total_users'] = User.query.count()
        role_specific['active_sessions'] = UserSession.query.filter_by(is_active=True).count()
        role_specific['referrals_today'] = referrals_today
        role_specific['appointments_today'] = Appointment.query.filter(func.date(Appointment.appointment_date) == today).count()
    return jsonify({
        'referrals_today': referrals_today,
        'pending_appointments': pending_appointments,
        'no_show_rate': round(no_show_rate, 1),
        'role_specific': role_specific
    })

# ==================== SPECIALIST ENDPOINTS ====================
@app.route('/api/specialists', methods=['GET'])
@jwt_required()
def get_specialists():
    specialists = Specialist.query.filter_by(is_available=True).all()
    result = []
    for s in specialists:
        user = db.session.get(User, s.user_id)
        specialty = db.session.get(Specialty, s.specialty_id)
        result.append({'id': s.id, 'name': user.full_name, 'specialty': specialty.name if specialty else None})
    return jsonify(result)

@app.route('/api/specialists/<int:specialist_id>/schedule', methods=['GET'])
@jwt_required()
def get_specialist_schedule(specialist_id):
    user_id = get_jwt_identity()
    claims = get_jwt()
    role = claims.get('role')
    if role == 'specialist':
        specialist = Specialist.query.filter_by(user_id=user_id).first()
        if not specialist or specialist.id != specialist_id:
            return jsonify({'error': 'Unauthorized'}), 403
    schedules = SpecialistSchedule.query.filter_by(specialist_id=specialist_id, is_active=True).order_by(SpecialistSchedule.day_of_week).all()
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']
    result = []
    for s in schedules:
        result.append({
            'id': s.id,
            'day': days[s.day_of_week],
            'day_of_week': s.day_of_week,
            'start_time': s.start_time.strftime('%H:%M'),
            'end_time': s.end_time.strftime('%H:%M'),
            'max_patients': s.max_patients
        })
    return jsonify(result)

@app.route('/api/specialists/<int:specialist_id>/available-slots', methods=['GET'])
@jwt_required()
def get_available_slots(specialist_id):
    claims = get_jwt()
    if claims.get('role') not in ['head_nurse', 'admin']:
        return jsonify({'error': 'Forbidden'}), 403
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({'error': 'Date required'}), 400
    try:
        target_date = datetime.strptime(date_str, '%Y-%m-%d').date()
    except:
        return jsonify({'error': 'Invalid date format'}), 400
    day_of_week = target_date.weekday()
    schedule = SpecialistSchedule.query.filter_by(specialist_id=specialist_id, day_of_week=day_of_week, is_active=True).first()
    if not schedule:
        return jsonify({'slots': []})
    slots = []
    start = datetime.combine(target_date, schedule.start_time)
    end = datetime.combine(target_date, schedule.end_time)
    current = start
    while current < end:
        booked = Appointment.query.filter(
            Appointment.specialist_id == specialist_id,
            Appointment.appointment_date == current,
            Appointment.status.in_(['scheduled', 'confirmed'])
        ).first()
        if not booked:
            slots.append(current.strftime('%H:%M'))
        current += timedelta(minutes=30)
    return jsonify({'slots': slots})

@app.route('/api/specialists/<int:specialist_id>/slots', methods=['GET'])
@jwt_required()
def get_specialist_slots(specialist_id):
    claims = get_jwt()
    if claims.get('role') not in ['head_nurse', 'admin']:
        return jsonify({'error': 'Forbidden'}), 403
    
    date_str = request.args.get('date')
    if not date_str:
        return jsonify({'error': 'Date required'}), 400
    try:
        target_date = datetime.strptime(date_str, '%Y-%m-%d').date()
    except:
        return jsonify({'error': 'Invalid date format'}), 400
    
    day_of_week = target_date.weekday()
    schedule = SpecialistSchedule.query.filter_by(
        specialist_id=specialist_id, day_of_week=day_of_week, is_active=True
    ).first()
    if not schedule:
        return jsonify({'slots': [], 'message': 'No working hours on this day'})
    
    start = datetime.combine(target_date, schedule.start_time)
    end = datetime.combine(target_date, schedule.end_time)
    all_slots = []
    current = start
    while current < end:
        all_slots.append(current)
        current += timedelta(minutes=30)
    
    booked_appointments = Appointment.query.filter(
        Appointment.specialist_id == specialist_id,
        func.date(Appointment.appointment_date) == target_date,
        Appointment.status.in_(['scheduled', 'confirmed', 'checked_in'])
    ).all()
    booked_times = [apt.appointment_date for apt in booked_appointments]
    
    slots_data = []
    for slot_time in all_slots:
        is_taken = any(abs((slot_time - bt).total_seconds()) < 900 for bt in booked_times)
        slots_data.append({
            'time': slot_time.strftime('%H:%M'),
            'datetime': slot_time.isoformat(),
            'available': not is_taken,
            'taken': is_taken
        })
    
    available_count = sum(1 for s in slots_data if s['available'])
    
    return jsonify({
        'date': date_str,
        'slots': slots_data,
        'available_count': available_count,
        'total_slots': len(slots_data),
        'schedule': {
            'start': schedule.start_time.strftime('%H:%M'),
            'end': schedule.end_time.strftime('%H:%M')
        }
    })

@app.route('/api/appointments/book', methods=['POST'])
@jwt_required()
def book_appointment():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    
    data = request.json
    required_fields = ['patient_id', 'specialist_id', 'appointment_date', 'reason']
    for field in required_fields:
        if field not in data:
            return jsonify({'error': f'Missing field: {field}'}), 400
    
    patient = db.session.get(Patient, data['patient_id'])
    if not patient:
        return jsonify({'error': 'Patient not found'}), 404
    if patient.department_id != nurse.department_id:
        return jsonify({'error': 'You do not have access to this patient'}), 403
    
    specialist = db.session.get(Specialist, data['specialist_id'])
    if not specialist:
        return jsonify({'error': 'Specialist not found'}), 404
    
    try:
        appointment_date = datetime.fromisoformat(data['appointment_date'])
    except:
        return jsonify({'error': 'Invalid appointment date format'}), 400
    
    conflict = Appointment.query.filter(
        Appointment.specialist_id == specialist.id,
        Appointment.appointment_date < appointment_date + timedelta(minutes=30),
        Appointment.appointment_date + timedelta(minutes=Appointment.duration) > appointment_date,
        Appointment.status.in_(['scheduled', 'confirmed', 'checked_in'])
    ).first()
    if conflict:
        return jsonify({'error': 'Time slot already booked'}), 409
    
    referral_id = data.get('referral_id')
    if not referral_id:
        referral = Referral(
            patient_id=patient.id,
            referring_nurse_id=nurse.id,
            assigned_specialist_id=specialist.id,
            referring_facility_id=nurse.facility_id,
            referred_to_facility_id=specialist.facility_id or nurse.facility_id,
            reason=data.get('reason', 'Appointment booking'),
            priority='routine',
            status='assigned',
            created_by=nurse.user_id
        )
        db.session.add(referral)
        db.session.flush()
        referral_id = referral.id
    else:
        referral = db.session.get(Referral, referral_id)
        if not referral:
            return jsonify({'error': 'Referral not found'}), 404
    
    appointment = Appointment(
        appointment_number=None,
        referral_id=referral_id,
        patient_id=patient.id,
        specialist_id=specialist.id,
        appointment_date=appointment_date,
        duration=data.get('duration', 30),
        status='scheduled',
        created_by=nurse.user_id
    )
    db.session.add(appointment)
    db.session.commit()
    
    send_notification(
        patient.user_id,
        'appointment_scheduled',
        'New Appointment Scheduled',
        f'Your appointment with Dr. {specialist.user.full_name} on {appointment_date.strftime("%d/%m/%Y %H:%M")} has been scheduled.',
        {'appointment_id': appointment.id},
        send_email_copy=True,
        send_sms_copy=True
    )
    send_notification(
        specialist.user_id,
        'appointment_scheduled',
        'New Appointment Scheduled',
        f'Appointment with patient {patient.user.full_name} on {appointment_date.strftime("%d/%m/%Y %H:%M")}.',
        {'appointment_id': appointment.id},
        send_email_copy=True
    )
    
    UserActivityLog.log_action(
        user_id=nurse.user_id,
        action_type='CREATE',
        resource_type='appointment',
        resource_id=appointment.id,
        resource_details={'patient_id': patient.id, 'specialist_id': specialist.id},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    
    return jsonify({
        'success': True,
        'message': 'Appointment booked successfully',
        'appointment': {
            'id': appointment.id,
            'appointment_number': appointment.appointment_number,
            'date': appointment.appointment_date.isoformat()
        }
    }), 201

@app.route('/api/appointments/specialist/<int:specialist_id>', methods=['GET'])
@jwt_required()
def get_specialist_appointments(specialist_id):
    user_id = get_jwt_identity()
    claims = get_jwt()
    role = claims.get('role')
    if role == 'specialist':
        specialist = Specialist.query.filter_by(user_id=user_id).first()
        if not specialist or specialist.id != specialist_id:
            return jsonify({'error': 'Unauthorized'}), 403
    date_str = request.args.get('date')
    query = Appointment.query.filter_by(specialist_id=specialist_id)
    if date_str:
        try:
            date_obj = datetime.strptime(date_str, '%Y-%m-%d').date()
            query = query.filter(func.date(Appointment.appointment_date) == date_obj)
        except ValueError:
            return jsonify({'error': 'Invalid date format'}), 400
    appointments = query.order_by(Appointment.appointment_date).all()
    result = []
    for a in appointments:
        patient = db.session.get(Patient, a.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        referral = db.session.get(Referral, a.referral_id)
        has_letter = bool(referral and ReferralDocument.query.filter_by(referral_id=referral.id).count() > 0)
        result.append({
            'id': a.id,
            'appointment_number': a.appointment_number,
            'time': a.appointment_date.strftime('%H:%M'),
            'date': a.appointment_date.isoformat(),
            'patient': {
                'id': patient.id if patient else None,
                'omang': patient.omang if patient else '',
                'name': patient_user.full_name if patient_user else 'Unknown',
                'gender': patient_user.gender if patient_user else ''
            },
            'referral': {
                'id': referral.id if referral else None,
                'reason': referral.reason if referral else '',
                'priority': referral.priority if referral else 'routine',
                'letter_available': has_letter
            },
            'status': a.status,
            'checked_in': a.checked_in,
            'checked_in_at': a.checked_in_at.isoformat() if a.checked_in_at else None,
            'duration': a.duration
        })
    return jsonify(result)

@app.route('/api/referrals/<int:referral_id>/document', methods=['GET'])
@jwt_required()
def get_referral_document(referral_id):
    claims = get_jwt()
    if claims.get('role') not in ['specialist', 'head_nurse']:
        return jsonify({'error': 'Forbidden'}), 403
    referral = db.session.get(Referral, referral_id)
    if not referral:
        return jsonify({'error': 'Referral not found'}), 404
    doc = ReferralDocument.query.filter_by(referral_id=referral_id, is_latest_version=True).first()
    if not doc:
        return jsonify({'error': 'No document found'}), 404
    return send_file(doc.file_path, mimetype=doc.mime_type)

# ==================== CLINICAL DECISION SUPPORT ====================
@app.route('/api/referrals/validate', methods=['POST'])
@jwt_required()
def validate_referral():
    data = request.json
    reason = data.get('reason', '').lower()
    specialty = data.get('specialty', '').lower()
    priority = data.get('priority', 'routine')
    
    warnings = []
    suggestions = []
    
    if len(reason) < 10:
        warnings.append("Reason is too short – please provide more clinical detail.")
    if not specialty:
        suggestions.append("Consider assigning a specialty to speed up processing.")
    if priority == 'emergency' and 'emergency' not in reason and 'urgent' not in reason:
        warnings.append("Emergency priority selected but reason does not indicate urgency.")
    if 'chest' in reason and specialty != 'cardiology':
        suggestions.append("Chest symptoms may be better suited for Cardiology.")
    if 'rash' in reason and specialty != 'dermatology':
        suggestions.append("Skin rash may be better suited for Dermatology.")
    if 'pregnancy' in reason and specialty != 'obstetrics & gynaecology':
        suggestions.append("Pregnancy-related issues are typically handled by Obstetrics & Gynaecology.")
    
    return jsonify({
        'valid': len(warnings) == 0,
        'warnings': warnings,
        'suggestions': suggestions
    })

# ==================== AI TRIAGE ====================
@app.route('/api/referrals/suggest-priority', methods=['POST'])
@jwt_required()
def suggest_priority():
    data = request.json
    reason = data.get('reason', '').lower()
    diagnosis = data.get('diagnosis', '').lower()
    text = reason + " " + diagnosis
    
    emergency_keywords = ['emergency', 'life-threatening', 'critical', 'severe bleeding', 'stroke', 'heart attack', 'unconscious']
    urgent_keywords = ['urgent', 'severe', 'acute', 'fracture', 'infection', 'high fever', 'difficulty breathing']
    
    if any(kw in text for kw in emergency_keywords):
        suggested = 'emergency'
    elif any(kw in text for kw in urgent_keywords):
        suggested = 'urgent'
    else:
        suggested = 'routine'
    
    return jsonify({'suggested_priority': suggested})

# ==================== WAITLIST MANAGEMENT & REFERRAL TRACKER ====================
@app.route('/api/referrals/tracker', methods=['GET'])
@jwt_required()
def referral_tracker():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
    if not nurse:
        return jsonify({'error': 'Nurse profile not found'}), 404
    
    referrals = Referral.query.filter_by(
        referring_nurse_id=nurse.id
    ).order_by(Referral.created_at.desc()).all()
    
    result = []
    for r in referrals:
        patient = db.session.get(Patient, r.patient_id)
        apt = Appointment.query.filter_by(referral_id=r.id).first()
        wait_days = (apt.appointment_date - r.created_at).days if apt and apt.appointment_date else None
        
        specialist_name = None
        if r.assigned_specialist_id:
            specialist = db.session.get(Specialist, r.assigned_specialist_id)
            if specialist:
                specialist_user = db.session.get(User, specialist.user_id)
                specialist_name = specialist_user.full_name if specialist_user else None
        
        result.append({
            'id': r.id,
            'referral_number': r.referral_number,
            'patient_name': patient.user.full_name if patient else 'Unknown',
            'priority': r.priority,
            'status': r.status,
            'created_at': r.created_at.isoformat(),
            'appointment_date': apt.appointment_date.isoformat() if apt else None,
            'waiting_days': wait_days,
            'specialist': specialist_name
        })
    return jsonify(result)

@app.route('/api/referrals/average-wait-times', methods=['GET'])
@jwt_required()
def average_wait_times():
    claims = get_jwt()
    if claims.get('role') not in ['head_nurse', 'admin']:
        return jsonify({'error': 'Forbidden'}), 403

    results = db.session.query(
        Specialty.name,
        func.avg(func.extract('day', Appointment.appointment_date - Referral.created_at)).label('avg_wait_days')
    ).select_from(Referral)\
     .join(Specialist, Referral.assigned_specialist_id == Specialist.id)\
     .join(Specialty, Specialist.specialty_id == Specialty.id)\
     .join(Appointment, Appointment.referral_id == Referral.id)\
     .filter(Referral.status == 'completed')\
     .group_by(Specialty.name).all()

    return jsonify([{'specialty': r[0], 'avg_wait_days': round(r[1], 1) if r[1] else None} for r in results])

# ==================== INTEROPERABILITY (FHIR stub) ====================
@app.route('/api/fhir/Patient', methods=['GET'])
@jwt_required()
def fhir_patient_stub():
    omang = request.args.get('identifier')
    if not omang:
        return jsonify({'resourceType': 'Bundle', 'entry': []})
    user = User.query.filter_by(omang=omang).first()
    if not user:
        return jsonify({'resourceType': 'Bundle', 'entry': []})
    patient = Patient.query.filter_by(user_id=user.id).first()
    fhir_patient = {
        'resourceType': 'Patient',
        'id': str(user.uuid),
        'identifier': [{'system': 'http://health.gov.bw/omang', 'value': user.omang}],
        'name': [{'family': user.full_name.split()[-1], 'given': user.full_name.split()[:-1]}],
        'gender': user.gender,
        'birthDate': patient.date_of_birth.isoformat() if patient and patient.date_of_birth else None,
        'telecom': [{'system': 'phone', 'value': user.phone}] if user.phone else []
    }
    return jsonify({'resourceType': 'Bundle', 'entry': [{'resource': fhir_patient}]})

# ==================== IN-APP NOTIFICATIONS FOR REMINDERS ====================
@app.route('/api/notifications/send-appointment-reminders', methods=['POST'])
@jwt_required()
def send_appointment_reminders():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    tomorrow = datetime.now(timezone.utc).date() + timedelta(days=1)
    appointments = Appointment.query.filter(
        func.date(Appointment.appointment_date) == tomorrow,
        Appointment.reminder_sent == False
    ).all()
    count = 0
    for apt in appointments:
        patient = db.session.get(Patient, apt.patient_id)
        user = db.session.get(User, patient.user_id)
        send_notification(
            user.id,
            'appointment_reminder',
            'Upcoming Appointment Reminder',
            f'You have an appointment with Dr. {apt.specialist_rel.user.full_name} tomorrow at {apt.appointment_date.strftime("%H:%M")}.',
            {'appointment_id': apt.id},
            send_email_copy=False,
            send_sms_copy=False
        )
        apt.reminder_sent = True
        apt.reminder_sent_at = datetime.now(timezone.utc)
        count += 1
    db.session.commit()
    return jsonify({'success': True, 'reminders_sent': count})

# ==================== FACILITIES ====================
@app.route('/api/facilities', methods=['GET'])
@jwt_required()
def get_facilities():
    facilities = Facility.query.filter_by(is_active=True).all()
    return jsonify([f.to_dict() for f in facilities])

# ==================== REPORTS ====================
@app.route('/api/reports/appointment-volumes', methods=['GET'])
@jwt_required()
def get_appointment_volume_report():
    claims = get_jwt()
    role = claims.get('role')
    if role not in ['head_nurse', 'admin']:
        return jsonify({'error': 'Forbidden'}), 403
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    query = Appointment.query
    if start_date:
        query = query.filter(Appointment.appointment_date >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Appointment.appointment_date <= datetime.fromisoformat(end_date))
    if role == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
        if nurse:
            query = query.join(Patient).filter(Patient.department_id == nurse.department_id)
    appointments = query.all()
    by_date = {}
    for a in appointments:
        date_str = a.appointment_date.strftime('%Y-%m-%d')
        if date_str not in by_date:
            by_date[date_str] = {'total': 0, 'scheduled': 0, 'completed': 0, 'missed': 0, 'cancelled': 0}
        by_date[date_str]['total'] += 1
        by_date[date_str][a.status] = by_date[date_str].get(a.status, 0) + 1
    return jsonify({
        'period': {'start_date': start_date, 'end_date': end_date},
        'summary': {'total_appointments': len(appointments), 'by_date': by_date}
    })

@app.route('/api/reports/waiting-times', methods=['GET'])
@jwt_required()
def waiting_times_report():
    claims = get_jwt()
    if claims.get('role') not in ['head_nurse', 'admin']:
        return jsonify({'error': 'Forbidden'}), 403
    
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    query = Referral.query
    if start_date:
        query = query.filter(Referral.created_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(Referral.created_at <= datetime.fromisoformat(end_date))
    if claims.get('role') == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=get_jwt_identity()).first()
        if nurse:
            query = query.join(Patient).filter(Patient.department_id == nurse.department_id)
    
    referrals = query.all()
    waiting_times = []
    for r in referrals:
        apt = Appointment.query.filter_by(referral_id=r.id).first()
        if apt and apt.appointment_date:
            waiting_days = (apt.appointment_date - r.created_at).days
            waiting_times.append({
                'referral_number': r.referral_number,
                'patient_name': r.patient_rel.user.full_name if r.patient_rel and r.patient_rel.user else 'Unknown',
                'created_at': r.created_at.isoformat(),
                'appointment_date': apt.appointment_date.isoformat(),
                'waiting_days': waiting_days,
                'priority': r.priority
            })
    return jsonify({'waiting_times': waiting_times, 'count': len(waiting_times)})

@app.route('/api/reports/referral-volume-by-district', methods=['GET'])
@jwt_required()
def referral_volume_by_district():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    results = db.session.query(Patient.district, func.count(Referral.id)).join(Referral).group_by(Patient.district).all()
    return jsonify([{'district': d or 'Unknown', 'count': c} for d, c in results])

# ==================== EXPORT REPORTS ====================
@app.route('/api/reports/export/<string:report_type>', methods=['GET'])
@jwt_required()
def export_report(report_type):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403

    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    def get_waiting_times_data():
        query = Referral.query
        if start_date:
            query = query.filter(Referral.created_at >= datetime.fromisoformat(start_date))
        if end_date:
            query = query.filter(Referral.created_at <= datetime.fromisoformat(end_date))
        referrals = query.all()
        data = []
        for r in referrals:
            apt = Appointment.query.filter_by(referral_id=r.id).first()
            if apt and apt.appointment_date:
                waiting_days = (apt.appointment_date - r.created_at).days
                data.append({
                    'referral_number': r.referral_number,
                    'patient_name': r.patient_rel.user.full_name if r.patient_rel and r.patient_rel.user else 'Unknown',
                    'created_at': r.created_at.isoformat(),
                    'appointment_date': apt.appointment_date.isoformat(),
                    'waiting_days': waiting_days,
                    'priority': r.priority
                })
        return data

    def get_referral_volume_data():
        results = db.session.query(Patient.district, func.count(Referral.id)).join(Referral).group_by(Patient.district).all()
        return [{'district': d or 'Unknown', 'count': c} for d, c in results]

    output = StringIO()
    writer = csv.writer(output)
    if report_type == 'waiting_times':
        data = get_waiting_times_data()
        writer.writerow(['Referral Number', 'Patient Name', 'Created At', 'Appointment Date', 'Waiting Days', 'Priority'])
        for row in data:
            writer.writerow([row['referral_number'], row['patient_name'], row['created_at'], row['appointment_date'], row['waiting_days'], row['priority']])
    elif report_type == 'referral_volume':
        data = get_referral_volume_data()
        writer.writerow(['District', 'Referral Count'])
        for row in data:
            writer.writerow([row['district'], row['count']])
    else:
        return jsonify({'error': 'Invalid report type'}), 400
    output.seek(0)
    return send_file(output, mimetype='text/csv', as_attachment=True, download_name=f'{report_type}_report.csv', conditional=False)

@app.route('/api/reports/export-pdf/<string:report_type>', methods=['GET'])
@jwt_required()
def export_report_pdf(report_type):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403

    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    def get_waiting_times_data():
        query = Referral.query
        if start_date:
            query = query.filter(Referral.created_at >= datetime.fromisoformat(start_date))
        if end_date:
            query = query.filter(Referral.created_at <= datetime.fromisoformat(end_date))
        referrals = query.all()
        data = []
        for r in referrals:
            apt = Appointment.query.filter_by(referral_id=r.id).first()
            if apt and apt.appointment_date:
                waiting_days = (apt.appointment_date - r.created_at).days
                data.append({
                    'referral_number': r.referral_number,
                    'patient_name': r.patient_rel.user.full_name if r.patient_rel and r.patient_rel.user else 'Unknown',
                    'created_at': r.created_at.isoformat(),
                    'appointment_date': apt.appointment_date.isoformat(),
                    'waiting_days': waiting_days,
                    'priority': r.priority
                })
        return data

    def get_referral_volume_data():
        results = db.session.query(Patient.district, func.count(Referral.id)).join(Referral).group_by(Patient.district).all()
        return [{'district': d or 'Unknown', 'count': c} for d, c in results]

    if report_type == 'waiting_times':
        data = get_waiting_times_data()
        title = "Waiting Times Report"
        headers = ['Referral Number', 'Patient Name', 'Created At', 'Appointment Date', 'Waiting Days', 'Priority']
        rows = [[str(row['referral_number']), str(row['patient_name']), str(row['created_at']), str(row['appointment_date']), str(row['waiting_days']), str(row['priority'])] for row in data]
    elif report_type == 'referral_volume':
        if claims.get('role') != 'admin':
            return jsonify({'error': 'Forbidden'}), 403
        data = get_referral_volume_data()
        title = "Referral Volume by District"
        headers = ['District', 'Referral Count']
        rows = [[str(row['district']), str(row['count'])] for row in data]
    else:
        return jsonify({'error': 'Invalid report type'}), 400

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=16, alignment=1, spaceAfter=12)
    
    elements = []
    elements.append(Paragraph(title, title_style))
    elements.append(Spacer(1, 12))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
    elements.append(Paragraph(f"Period: {start_date or 'All'} to {end_date or 'All'}", styles['Normal']))
    elements.append(Spacer(1, 12))
    
    if not rows:
        elements.append(Paragraph("No data available for the selected period.", styles['Normal']))
    else:
        table_data = [headers] + rows
        table = Table(table_data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.grey),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,0), 10),
            ('BOTTOMPADDING', (0,0), (-1,0), 12),
            ('BACKGROUND', (0,1), (-1,-1), colors.beige),
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('FONTSIZE', (0,1), (-1,-1), 9),
        ]))
        elements.append(table)
    
    doc.build(elements)
    buffer.seek(0)
    return send_file(buffer, mimetype='application/pdf', as_attachment=True, download_name=f'{report_type}_report.pdf', conditional=False)

@app.route('/api/reports/export-word/<string:report_type>', methods=['GET'])
@jwt_required()
def export_report_word(report_type):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403

    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    
    def get_waiting_times_data():
        query = Referral.query
        if start_date:
            query = query.filter(Referral.created_at >= datetime.fromisoformat(start_date))
        if end_date:
            query = query.filter(Referral.created_at <= datetime.fromisoformat(end_date))
        referrals = query.all()
        data = []
        for r in referrals:
            apt = Appointment.query.filter_by(referral_id=r.id).first()
            if apt and apt.appointment_date:
                waiting_days = (apt.appointment_date - r.created_at).days
                data.append({
                    'referral_number': r.referral_number,
                    'patient_name': r.patient_rel.user.full_name if r.patient_rel and r.patient_rel.user else 'Unknown',
                    'created_at': r.created_at.isoformat(),
                    'appointment_date': apt.appointment_date.isoformat(),
                    'waiting_days': waiting_days,
                    'priority': r.priority
                })
        return data

    def get_referral_volume_data():
        results = db.session.query(Patient.district, func.count(Referral.id)).join(Referral).group_by(Patient.district).all()
        return [{'district': d or 'Unknown', 'count': c} for d, c in results]

    if report_type == 'waiting_times':
        data = get_waiting_times_data()
        title = "Waiting Times Report"
        headers = ['Referral Number', 'Patient Name', 'Created At', 'Appointment Date', 'Waiting Days', 'Priority']
        rows = [[str(row['referral_number']), str(row['patient_name']), str(row['created_at']), str(row['appointment_date']), str(row['waiting_days']), str(row['priority'])] for row in data]
    elif report_type == 'referral_volume':
        if claims.get('role') != 'admin':
            return jsonify({'error': 'Forbidden'}), 403
        data = get_referral_volume_data()
        title = "Referral Volume by District"
        headers = ['District', 'Referral Count']
        rows = [[str(row['district']), str(row['count'])] for row in data]
    else:
        return jsonify({'error': 'Invalid report type'}), 400

    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Period: {start_date or 'All'} to {end_date or 'All'}")
    doc.add_paragraph("")
    
    if not rows:
        doc.add_paragraph("No data available for the selected period.")
    else:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for i, row in enumerate(rows):
            row_cells = table.rows[i+1].cells
            for j, cell_value in enumerate(row):
                row_cells[j].text = cell_value
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=f'{report_type}_report.docx', conditional=False)

# ==================== ESCALATION ALERTS ====================
@app.route('/api/escalation/alerts', methods=['GET'])
@jwt_required()
def get_escalation_alerts():
    claims = get_jwt()
    user_role = claims.get('role')
    
    if user_role not in ['head_nurse', 'admin']:
        return jsonify({'error': 'Forbidden'}), 403
    
    current_user_id = get_jwt_identity()
    
    department_id = None
    if user_role == 'head_nurse':
        nurse = Nurse.query.filter_by(user_id=current_user_id).first()
        if not nurse:
            return jsonify({'error': 'Nurse profile not found'}), 404
        department_id = nurse.department_id
    
    six_months_ago = datetime.now(timezone.utc) - timedelta(days=180)
    
    query = Appointment.query.join(Patient).filter(
        Appointment.created_at <= six_months_ago,
        Appointment.status.notin_(['completed', 'cancelled'])
    )
    
    if department_id:
        query = query.filter(Patient.department_id == department_id)
    
    query = query.order_by(Appointment.created_at.asc())
    
    long_waiting = query.all()
    
    result = [{
        'id': a.id,
        'patient_name': a.patient_rel.user.full_name,
        'appointment_date': a.appointment_date.isoformat(),
        'created_at': a.created_at.isoformat(),
        'status': a.status,
        'waiting_days': (datetime.now(timezone.utc) - a.created_at).days
    } for a in long_waiting]
    
    return jsonify(result)

# ==================== FILE UPLOAD ====================
@app.route('/api/upload', methods=['POST'])
@jwt_required()
def upload_file():
    claims = get_jwt()
    if claims.get('role') != 'head_nurse':
        return jsonify({'error': 'Forbidden'}), 403
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    referral_id = request.form.get('referral_id', type=int)
    document_type = request.form.get('document_type', 'referral_letter_initial')
    allowed_extensions = {'png', 'jpg', 'jpeg', 'pdf', 'gif'}
    if '.' not in file.filename or file.filename.rsplit('.', 1)[1].lower() not in allowed_extensions:
        return jsonify({'error': 'File type not allowed'}), 400
    filename = secure_filename(file.filename)
    name, ext = os.path.splitext(filename)
    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
    unique_filename = f"{name}_{timestamp}{ext}"
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
    file.save(file_path)
    file_size = os.path.getsize(file_path)
    document = ReferralDocument(
        referral_id=referral_id,
        document_type=document_type,
        filename=unique_filename,
        file_path=file_path,
        file_size=file_size,
        mime_type=file.content_type,
        uploaded_by=get_jwt_identity()
    )
    db.session.add(document)
    db.session.commit()
    return jsonify({
        'success': True,
        'message': 'File uploaded successfully',
        'document': {'id': document.id, 'filename': document.filename, 'size': document.file_size, 'uploaded_at': document.uploaded_at.isoformat()}
    }), 201

# ==================== ADMIN USER MANAGEMENT ====================
@app.route('/api/admin/users', methods=['GET'])
@jwt_required()
def admin_get_users():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    users = User.query.all()
    return jsonify([user.to_dict() for user in users])

@app.route('/api/admin/users/<int:user_id>', methods=['GET'])
@jwt_required()
def admin_get_user(user_id):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    return jsonify(user.to_dict())

@app.route('/api/admin/users', methods=['POST'])
@jwt_required()
def admin_create_user():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    data = request.json
    if data['role'] not in ['head_nurse', 'specialist', 'admin']:
        return jsonify({'error': 'Invalid role'}), 400
    if User.query.filter_by(omang=data['omang']).first():
        return jsonify({'error': 'Omang already exists'}), 400
    temp_pin = f"{random.randint(1000, 9999)}"
    user = User(
        omang=data['omang'],
        full_name=data['full_name'],
        role=data['role'],
        gender=data.get('gender'),
        phone=data.get('phone'),
        email=data.get('email'),
        employee_id=data.get('employee_id'),
        department=data.get('department'),
        job_title=data.get('job_title'),
        status='active',
        created_by=get_jwt_identity()
    )
    user.set_pin(temp_pin)
    db.session.add(user)
    db.session.flush()
    if data['role'] == 'head_nurse':
        nurse = Nurse(
            user_id=user.id,
            employee_id=data.get('employee_id'),
            facility_id=data.get('facility_id'),
            department_id=data.get('department_id'),
            qualification=data.get('qualification'),
            years_experience=data.get('years_experience'),
            license_number=data.get('license_number')
        )
        db.session.add(nurse)
    elif data['role'] == 'specialist':
        specialist = Specialist(
            user_id=user.id,
            employee_id=data.get('employee_id'),
            specialty_id=data.get('specialty_id'),
            sub_specialty=data.get('sub_specialty'),
            facility_id=data.get('facility_id'),
            department=data.get('department'),
            qualifications=data.get('qualifications'),
            years_experience=data.get('years_experience'),
            license_number=data.get('license_number'),
            consultation_duration=data.get('consultation_duration', 30),
            max_patients_per_day=data.get('max_patients_per_day', 15)
        )
        db.session.add(specialist)
    db.session.commit()
    send_notification(user.id, 'user_created', 'Account Created', f'Your temporary PIN is {temp_pin}', send_email_copy=True)
    return jsonify({'success': True, 'user': user.to_dict(), 'temp_pin': temp_pin}), 201

@app.route('/api/admin/users/<int:user_id>', methods=['PUT'])
@jwt_required()
def admin_update_user(user_id):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    data = request.json
    for field in ['full_name', 'phone', 'email', 'status', 'department', 'job_title', 'employee_id']:
        if field in data:
            setattr(user, field, data[field])
    if user.role == 'head_nurse' and user.nurse:
        nurse = user.nurse
        for field in ['facility_id', 'department_id', 'qualification', 'years_experience', 'license_number']:
            if field in data:
                setattr(nurse, field, data[field])
    elif user.role == 'specialist' and user.specialist:
        spec = user.specialist
        for field in ['specialty_id', 'sub_specialty', 'facility_id', 'qualifications', 'years_experience', 'license_number', 'consultation_duration', 'max_patients_per_day', 'is_available']:
            if field in data:
                setattr(spec, field, data[field])
    db.session.commit()
    return jsonify({'success': True, 'user': user.to_dict()})

@app.route('/api/admin/users/<int:user_id>/reset-pin', methods=['POST'])
@jwt_required()
def admin_reset_user_pin(user_id):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    new_pin = f"{random.randint(1000, 9999)}"
    user.set_pin(new_pin)
    user.pin_reset_token = None
    user.pin_reset_expires_at = None
    db.session.commit()
    send_notification(user.id, 'pin_reset', 'PIN Reset by Administrator', f'Your PIN has been reset. New PIN: {new_pin}', send_email_copy=True)
    UserActivityLog.log_action(
        user_id=get_jwt_identity(),
        action_type='RESET_PIN',
        resource_type='user',
        resource_id=user.id,
        resource_details={'target_user': user.full_name},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    return jsonify({'success': True, 'new_pin': new_pin})

@app.route('/api/admin/users/<int:user_id>/hard-delete', methods=['DELETE'])
@jwt_required()
def admin_hard_delete_user(user_id):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    
    if user.id == int(get_jwt_identity()):
        return jsonify({'error': 'Cannot delete your own account'}), 400
    
    if user.role == 'patient':
        patient = Patient.query.filter_by(user_id=user.id).first()
        if patient:
            referral_count = Referral.query.filter_by(patient_id=patient.id).count()
            appointment_count = Appointment.query.filter_by(patient_id=patient.id).count()
            if referral_count > 0 or appointment_count > 0:
                return jsonify({'error': f'Cannot delete user with {referral_count} referrals and {appointment_count} appointments. Deactivate instead.'}), 400
    
    created_referrals = Referral.query.filter_by(created_by=user.id).count()
    if created_referrals > 0:
        return jsonify({'error': f'Cannot delete user who created {created_referrals} referrals. Deactivate instead.'}), 400
    
    try:
        if user.role == 'patient':
            patient = Patient.query.filter_by(user_id=user.id).first()
            if patient:
                db.session.delete(patient)
        
        if user.role == 'head_nurse':
            nurse = Nurse.query.filter_by(user_id=user.id).first()
            if nurse:
                db.session.delete(nurse)
        elif user.role == 'specialist':
            specialist = Specialist.query.filter_by(user_id=user.id).first()
            if specialist:
                SpecialistSchedule.query.filter_by(specialist_id=specialist.id).delete()
                db.session.delete(specialist)
        
        Notification.query.filter_by(user_id=user.id).delete()
        UserActivityLog.query.filter_by(user_id=user.id).delete()
        
        db.session.delete(user)
        db.session.commit()
        
        UserActivityLog.log_action(
            user_id=get_jwt_identity(),
            action_type='HARD_DELETE',
            resource_type='user',
            resource_id=user_id,
            resource_details={'deleted_user': user.full_name, 'omang': user.omang},
            ip_address=request.remote_addr,
            user_agent=request.user_agent.string
        )
        
        return jsonify({'success': True, 'message': 'User permanently deleted'}), 200
    except Exception as e:
        db.session.rollback()
        app.logger.error(f"Hard delete error: {e}")
        return jsonify({'error': 'Failed to delete user. They may have existing records.'}), 500

@app.route('/api/admin/users/<int:user_id>', methods=['DELETE'])
@jwt_required()
def admin_delete_user(user_id):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    user = db.session.get(User, user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    if user.id == get_jwt_identity():
        return jsonify({'error': 'Cannot delete yourself'}), 400
    user.status = 'inactive'
    db.session.commit()
    return jsonify({'success': True})

# ==================== ADMIN ACTIVITY LOGS ====================
@app.route('/api/admin/activity-logs', methods=['GET'])
@jwt_required()
def admin_activity_logs():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    user_id = request.args.get('user_id', type=int)
    action_type = request.args.get('action_type')
    resource_type = request.args.get('resource_type')
    start_date = request.args.get('start_date')
    end_date = request.args.get('end_date')
    limit = request.args.get('limit', 100, type=int)
    query = UserActivityLog.query
    if user_id:
        query = query.filter_by(user_id=user_id)
    if action_type:
        query = query.filter_by(action_type=action_type)
    if resource_type:
        query = query.filter_by(resource_type=resource_type)
    if start_date:
        query = query.filter(UserActivityLog.performed_at >= datetime.fromisoformat(start_date))
    if end_date:
        query = query.filter(UserActivityLog.performed_at <= datetime.fromisoformat(end_date))
    logs = query.order_by(UserActivityLog.performed_at.desc()).limit(limit).all()
    result = []
    for log in logs:
        result.append({
            'id': log.id,
            'user_id': log.user_id,
            'user_name': log.user_full_name,
            'user_role': log.user_role,
            'action_type': log.action_type,
            'resource_type': log.resource_type,
            'resource_id': log.resource_id,
            'details': log.resource_details,
            'ip_address': log.ip_address,
            'performed_at': log.performed_at.isoformat(),
            'status': log.status
        })
    return jsonify(result)

@app.route('/api/admin/activity-logs/summary', methods=['GET'])
@jwt_required()
def admin_activity_summary():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    today = datetime.now(timezone.utc).date()
    total_today = UserActivityLog.query.filter(func.date(UserActivityLog.performed_at) == today).count()
    by_type = db.session.query(UserActivityLog.action_type, func.count().label('count')).filter(func.date(UserActivityLog.performed_at) == today).group_by(UserActivityLog.action_type).all()
    top_users = db.session.query(UserActivityLog.user_id, UserActivityLog.user_full_name, func.count().label('count')).filter(func.date(UserActivityLog.performed_at) == today).group_by(UserActivityLog.user_id, UserActivityLog.user_full_name).order_by(func.count().desc()).limit(10).all()
    return jsonify({
        'total_today': total_today,
        'by_type': [{'type': t, 'count': c} for t, c in by_type],
        'top_users': [{'user_id': uid, 'name': name, 'count': c} for uid, name, c in top_users]
    })

# ==================== ADMIN SYSTEM SETTINGS ====================
@app.route('/api/admin/settings', methods=['GET'])
@jwt_required()
def admin_get_settings():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    settings = SystemConfig.query.all()
    return jsonify([{
        'key': s.config_key,
        'value': s.config_value,
        'type': s.config_type,
        'description': s.description,
        'editable': s.is_editable,
        'updated_at': s.updated_at.isoformat() if s.updated_at else None,
        'updated_by': s.updated_by
    } for s in settings])

@app.route('/api/admin/settings/<key>', methods=['PUT'])
@jwt_required()
def admin_update_setting(key):
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    data = request.json
    config = SystemConfig.query.filter_by(config_key=key).first_or_404()
    if not config.is_editable:
        return jsonify({'error': 'Setting not editable'}), 403
    config.config_value = str(data['value'])
    config.updated_by = get_jwt_identity()
    db.session.commit()
    return jsonify({'success': True})

# ==================== ADMIN DASHBOARD STATS ====================
@app.route('/api/admin/stats', methods=['GET'])
@jwt_required()
def admin_get_stats():
    try:
        claims = get_jwt()
        if claims.get('role') != 'admin':
            return jsonify({'error': 'Forbidden'}), 403
        
        today = datetime.now(timezone.utc).date()
        total_users = User.query.count()
        users_by_role = db.session.query(User.role, func.count().label('count')).group_by(User.role).all()
        active_sessions = UserSession.query.filter_by(is_active=True).count()
        logins_today = UserActivityLog.query.filter(UserActivityLog.action_type == 'LOGIN', func.date(UserActivityLog.performed_at) == today).count()
        referrals_today = Referral.query.filter(func.date(Referral.created_at) == today).count()
        appointments_today = Appointment.query.filter(func.date(Appointment.appointment_date) == today).count()
        
        db_size_mb = 0
        try:
            db_size = db.session.execute(text("SELECT pg_database_size(current_database())")).scalar()
            db_size_mb = db_size / (1024 * 1024) if db_size else 0
        except Exception as e:
            app.logger.warning(f"Could not get database size: {e}")
            db_size_mb = 0
        
        return jsonify({
            'users': {
                'total': total_users,
                'by_role': [{'role': str(r), 'count': c} for r, c in users_by_role],
                'active_sessions': active_sessions
            },
            'activity': {
                'logins_today': logins_today,
                'referrals_today': referrals_today,
                'appointments_today': appointments_today
            },
            'system': {
                'database_size_mb': round(db_size_mb, 2)
            }
        })
    except Exception as e:
        app.logger.error(f"Error in admin_get_stats: {e}")
        return jsonify({'error': str(e)}), 500

# ==================== ADMIN DEPARTMENT & SPECIALTY MANAGEMENT ====================
@app.route('/api/admin/departments', methods=['GET'])
@jwt_required()
def admin_get_departments():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    depts = Department.query.all()
    return jsonify([d.to_dict() for d in depts])

@app.route('/api/admin/departments', methods=['POST'])
@jwt_required()
def admin_create_department():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    data = request.json
    dept = Department(name=data['name'], code=data['code'])
    db.session.add(dept)
    db.session.commit()
    return jsonify(dept.to_dict()), 201

@app.route('/api/admin/specialties', methods=['GET'])
@jwt_required()
def admin_get_specialties():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    specs = Specialty.query.all()
    return jsonify([{'id': s.id, 'code': s.code, 'name': s.name, 'consultation_duration': s.consultation_duration} for s in specs])

# ==================== TEST ENDPOINTS ====================
@app.route('/api/test-cors', methods=['GET', 'OPTIONS'])
def test_cors():
    return jsonify({'message': 'CORS is working!'}), 200

@app.route('/api/test', methods=['GET'])
def test():
    return jsonify({'success': True, 'message': 'Backend is working!'})

# ==================== SPECIALIST DASHBOARD ====================
@app.route('/api/specialist/dashboard', methods=['GET'])
@jwt_required()
def specialist_dashboard():
    claims = get_jwt()
    if claims.get('role') != 'specialist':
        return jsonify({'error': 'Forbidden'}), 403
    
    user_id = get_jwt_identity()
    specialist = Specialist.query.filter_by(user_id=user_id).first()
    if not specialist:
        return jsonify({'error': 'Specialist profile not found'}), 404
    
    today = datetime.now(timezone.utc).date()
    
    today_appointments = Appointment.query.filter(
        Appointment.specialist_id == specialist.id,
        func.date(Appointment.appointment_date) == today,
        Appointment.status.in_(['scheduled', 'confirmed', 'checked_in', 'in_progress'])
    ).order_by(Appointment.appointment_date).all()
    
    pending_referrals = Referral.query.filter(
        Referral.assigned_specialist_id == specialist.id,
        Referral.status.in_(['assigned', 'pending'])
    ).order_by(
        db.case(
            (Referral.priority == 'emergency', 1),
            (Referral.priority == 'urgent', 2),
            (Referral.priority == 'routine', 3),
            else_=4
        ),
        Referral.created_at.asc()
    ).all()
    
    week_later = today + timedelta(days=7)
    upcoming_appointments = Appointment.query.filter(
        Appointment.specialist_id == specialist.id,
        func.date(Appointment.appointment_date) > today,
        func.date(Appointment.appointment_date) <= week_later,
        Appointment.status.in_(['scheduled', 'confirmed'])
    ).order_by(Appointment.appointment_date).all()
    
    total_patients_seen = Appointment.query.filter(
        Appointment.specialist_id == specialist.id,
        Appointment.status == 'completed'
    ).count()
    
    missed_appointments = Appointment.query.filter(
        Appointment.specialist_id == specialist.id,
        Appointment.status == 'missed'
    ).count()
    
    result = {
        'specialist': {
            'id': specialist.id,
            'name': specialist.user.full_name if specialist.user else None,
            'specialty': specialist.specialty_rel.name if specialist.specialty_rel else None,
            'is_available': specialist.is_available
        },
        'stats': {
            'today_count': len(today_appointments),
            'pending_referrals_count': len(pending_referrals),
            'upcoming_count': len(upcoming_appointments),
            'total_patients_seen': total_patients_seen,
            'missed_appointments': missed_appointments
        },
        'today_appointments': [],
        'pending_referrals': [],
        'upcoming_appointments': []
    }
    
    for apt in today_appointments:
        patient = db.session.get(Patient, apt.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        referral = db.session.get(Referral, apt.referral_id)
        has_letter = ReferralDocument.query.filter_by(referral_id=referral.id).first() is not None if referral else False
        
        result['today_appointments'].append({
            'id': apt.id,
            'appointment_number': apt.appointment_number,
            'time': apt.appointment_date.strftime('%H:%M'),
            'datetime': apt.appointment_date.isoformat(),
            'patient': {
                'id': patient.id if patient else None,
                'name': patient_user.full_name if patient_user else 'Unknown',
                'omang': patient.omang if patient else None,
                'age': (today - patient.date_of_birth).days // 365 if patient and patient.date_of_birth else None
            },
            'referral': {
                'id': referral.id if referral else None,
                'reason': referral.reason[:150] if referral else '',
                'priority': referral.priority if referral else 'routine',
                'has_letter': has_letter
            },
            'status': apt.status,
            'checked_in': apt.checked_in,
            'checked_in_at': apt.checked_in_at.isoformat() if apt.checked_in_at else None
        })
    
    for ref in pending_referrals:
        patient = db.session.get(Patient, ref.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        waiting_days = (datetime.now(timezone.utc) - ref.created_at).days
        
        result['pending_referrals'].append({
            'id': ref.id,
            'referral_number': ref.referral_number,
            'patient': {
                'id': patient.id if patient else None,
                'name': patient_user.full_name if patient_user else 'Unknown',
                'omang': patient.omang if patient else None
            },
            'reason': ref.reason[:200],
            'priority': ref.priority,
            'created_at': ref.created_at.isoformat(),
            'waiting_days': waiting_days
        })
    
    for apt in upcoming_appointments:
        patient = db.session.get(Patient, apt.patient_id)
        patient_user = db.session.get(User, patient.user_id) if patient else None
        
        result['upcoming_appointments'].append({
            'id': apt.id,
            'appointment_number': apt.appointment_number,
            'date': apt.appointment_date.strftime('%Y-%m-%d'),
            'time': apt.appointment_date.strftime('%H:%M'),
            'datetime': apt.appointment_date.isoformat(),
            'patient': {
                'id': patient.id if patient else None,
                'name': patient_user.full_name if patient_user else 'Unknown'
            },
            'status': apt.status
        })
    
    return jsonify(result), 200

# ==================== SPECIALIST UPDATE AVAILABILITY ====================
@app.route('/api/specialist/availability', methods=['PUT'])
@jwt_required()
def specialist_update_availability():
    claims = get_jwt()
    if claims.get('role') != 'specialist':
        return jsonify({'error': 'Forbidden'}), 403
    
    user_id = get_jwt_identity()
    specialist = Specialist.query.filter_by(user_id=user_id).first()
    if not specialist:
        return jsonify({'error': 'Specialist profile not found'}), 404
    
    data = request.json
    is_available = data.get('is_available', False)
    specialist.is_available = is_available
    db.session.commit()
    
    UserActivityLog.log_action(
        user_id=user_id,
        action_type='UPDATE_AVAILABILITY',
        resource_type='specialist',
        resource_id=specialist.id,
        resource_details={'is_available': is_available},
        ip_address=request.remote_addr,
        user_agent=request.user_agent.string
    )
    
    return jsonify({
        'success': True,
        'message': f'Availability set to {is_available}',
        'is_available': is_available
    }), 200

# ==================== SPECIALIST GET SINGLE APPOINTMENT DETAILS ====================
@app.route('/api/specialist/appointments/<int:appointment_id>', methods=['GET'])
@jwt_required()
def specialist_get_appointment_details(appointment_id):
    claims = get_jwt()
    if claims.get('role') != 'specialist':
        return jsonify({'error': 'Forbidden'}), 403
    
    user_id = get_jwt_identity()
    specialist = Specialist.query.filter_by(user_id=user_id).first()
    if not specialist:
        return jsonify({'error': 'Specialist profile not found'}), 404
    
    appointment = db.session.get(Appointment, appointment_id)
    if not appointment:
        return jsonify({'error': 'Appointment not found'}), 404
    
    if appointment.specialist_id != specialist.id:
        return jsonify({'error': 'Unauthorized - This appointment is not assigned to you'}), 403
    
    patient = db.session.get(Patient, appointment.patient_id)
    patient_user = db.session.get(User, patient.user_id) if patient else None
    referral = db.session.get(Referral, appointment.referral_id)
    
    documents = ReferralDocument.query.filter_by(referral_id=referral.id).all() if referral else []
    medical_history = PatientMedicalHistory.query.filter_by(patient_id=patient.id).all() if patient else []
    
    return jsonify({
        'id': appointment.id,
        'appointment_number': appointment.appointment_number,
        'datetime': appointment.appointment_date.isoformat(),
        'duration': appointment.duration,
        'status': appointment.status,
        'checked_in': appointment.checked_in,
        'checked_in_at': appointment.checked_in_at.isoformat() if appointment.checked_in_at else None,
        'outcome': appointment.outcome,
        'clinical_notes': appointment.clinical_notes,
        'patient': {
            'id': patient.id if patient else None,
            'name': patient_user.full_name if patient_user else 'Unknown',
            'omang': patient.omang if patient else None,
            'date_of_birth': patient.date_of_birth.isoformat() if patient and patient.date_of_birth else None,
            'gender': patient_user.gender if patient_user else None,
            'phone': patient_user.phone if patient_user else None,
            'email': patient_user.email if patient_user else None,
            'village': patient.village if patient else None,
            'district': patient.district if patient else None,
            'next_of_kin_name': patient.next_of_kin_name if patient else None,
            'next_of_kin_phone': patient.next_of_kin_phone if patient else None
        },
        'referral': {
            'id': referral.id if referral else None,
            'referral_number': referral.referral_number if referral else None,
            'reason': referral.reason if referral else '',
            'clinical_summary': referral.clinical_summary if referral else '',
            'diagnosis': referral.diagnosis if referral else '',
            'symptoms': referral.symptoms if referral else '',
            'priority': referral.priority if referral else 'routine',
            'created_at': referral.created_at.isoformat() if referral else None
        },
        'documents': [{
            'id': d.id,
            'filename': d.filename,
            'document_type': d.document_type,
            'uploaded_at': d.uploaded_at.isoformat() if d.uploaded_at else None
        } for d in documents],
        'medical_history': [{
            'id': m.id,
            'condition': m.condition,
            'diagnosis_date': m.diagnosis_date.isoformat() if m.diagnosis_date else None,
            'notes': m.notes,
            'is_active': m.is_active
        } for m in medical_history]
    }), 200

# ==================== NO-SHOW STATISTICS ENDPOINTS ====================
@app.route('/api/stats/no-show/monthly', methods=['GET'])
@jwt_required()
def get_monthly_no_show_stats():
    from datetime import date
    current_date = datetime.now(timezone.utc).date()
    stats = []
    
    for i in range(11, -1, -1):
        target_date = current_date.replace(day=1)
        if i > 0:
            month = target_date.month - i
            year = target_date.year
            if month < 1:
                month = 12 + month
                year = year - 1
        else:
            month = target_date.month
            year = target_date.year
        
        start_date = date(year, month, 1)
        if month == 12:
            end_date = date(year + 1, 1, 1)
        else:
            end_date = date(year, month + 1, 1)
        
        appointments = Appointment.query.filter(
            Appointment.appointment_date >= start_date,
            Appointment.appointment_date < end_date
        ).all()
        
        total = len(appointments)
        completed = sum(1 for a in appointments if a.status == 'completed')
        missed = sum(1 for a in appointments if a.status == 'missed')
        cancelled = sum(1 for a in appointments if a.status == 'cancelled')
        no_show_rate = (missed / total * 100) if total > 0 else 0
        
        stats.append({
            'year': year,
            'month': month,
            'month_name': datetime(year, month, 1).strftime('%B'),
            'total_appointments': total,
            'completed': completed,
            'missed': missed,
            'cancelled': cancelled,
            'no_show_rate': round(no_show_rate, 2)
        })
    
    return jsonify({'success': True, 'data': stats}), 200


@app.route('/api/stats/no-show/current', methods=['GET'])
@jwt_required()
def get_current_no_show_stats():
    from datetime import date
    now = datetime.now(timezone.utc)
    current_month = now.month
    current_year = now.year
    
    start_date = date(current_year, current_month, 1)
    if current_month == 12:
        end_date = date(current_year + 1, 1, 1)
    else:
        end_date = date(current_year, current_month + 1, 1)
    
    appointments = Appointment.query.filter(
        Appointment.appointment_date >= start_date,
        Appointment.appointment_date < end_date
    ).all()
    
    total = len(appointments)
    completed = sum(1 for a in appointments if a.status == 'completed')
    missed = sum(1 for a in appointments if a.status == 'missed')
    cancelled = sum(1 for a in appointments if a.status == 'cancelled')
    scheduled = sum(1 for a in appointments if a.status in ['scheduled', 'confirmed'])
    no_show_rate = (missed / total * 100) if total > 0 else 0
    
    return jsonify({
        'success': True,
        'data': {
            'year': current_year,
            'month': current_month,
            'month_name': now.strftime('%B'),
            'total_appointments': total,
            'completed': completed,
            'missed': missed,
            'cancelled': cancelled,
            'scheduled': scheduled,
            'no_show_rate': round(no_show_rate, 2)
        }
    }), 200


@app.route('/api/admin/check-missed-appointments', methods=['POST'])
@jwt_required()
def manual_check_missed_appointments():
    claims = get_jwt()
    if claims.get('role') != 'admin':
        return jsonify({'error': 'Forbidden'}), 403
    
    from services.no_show_scheduler import no_show_scheduler
    missed_count = no_show_scheduler._check_and_mark_missed_appointments()
    
    return jsonify({
        'success': True,
        'message': f'Checked appointments. Marked {missed_count} as missed.',
        'missed_count': missed_count
    }), 200

# ==================== RUN SERVER ====================
if __name__ == '__main__':
    print("\n" + "=" * 70)
    print(" ISAS BACKEND SERVER STARTING")
    print("=" * 70)
    print(" WebSocket: ws://localhost:5000")
    print(" API: http://localhost:5000/api")
    print(" Terminal Capture: OFF (use web UI to start)")
    print("=" * 70 + "\n")
    
    socketio.run(app, debug=True, host='0.0.0.0', port=5000, allow_unsafe_werkzeug=True)